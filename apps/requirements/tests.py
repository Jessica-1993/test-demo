from unittest.mock import patch
from datetime import timedelta
from types import SimpleNamespace
import json
import tempfile

from django.contrib.auth.models import User
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from kombu.exceptions import OperationalError
from rest_framework.test import APIClient

from apps.configuration.models import LLMModelConfig, ProjectConfig, PromptConfig
from apps.project_knowledge.models import ProjectModule
from .integration import RequirementReviewService
from .models import RequirementContentBlock, RequirementDocument, RequirementFamily, RequirementImageAnalysis, RequirementIntegrationBatch, RequirementIntegrationDraft, RequirementIntegrationRun, RequirementItem, RequirementParseRun, RequirementRevision, RequirementVersion, TestCase as LibraryTestCase, TestCaseGenerationTask
from .services import DocumentExtractionService, LLMGeminiVisionService, LLMResponsesVisionService, QiniuStorageService, RequirementContextBuilder, RequirementImageAnalysisService, RequirementIntegrationService, StructuredRequirementParser, TestCaseGenerationError, TestCaseGenerationService
from .tasks import run_testcase_generation_task


class FakeUrlopenResponse:
    status = 200

    def __init__(self, payload, content_type="application/json"):
        self.payload = payload
        self.headers = {"Content-Type": content_type}

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, traceback):
        return False

    def read(self):
        if isinstance(self.payload, bytes):
            return self.payload
        return json.dumps(self.payload).encode("utf-8")


class DocumentExtractionTests(TestCase):
    def test_requirement_heading_sort_prefixes_are_removed(self):
        self.assertEqual(StructuredRequirementParser._clean_heading("2.1.2 补充说明"), "补充说明")
        self.assertEqual(StructuredRequirementParser._clean_heading("REQ-001 用户登录"), "用户登录")

    @override_settings(QINIU_DOMAIN="cdn.example.com/")
    def test_qiniu_domain_without_scheme_is_normalized(self):
        storage = QiniuStorageService()

        self.assertEqual(storage.domain, "https://cdn.example.com")
        self.assertEqual(storage.public_url("docs/需求.docx"), "https://cdn.example.com/docs/%E9%9C%80%E6%B1%82.docx")

    @override_settings(
        QINIU_DOMAIN="http://cdn.example.com",
        QINIU_ACCESS_KEY="access",
        QINIU_SECRET_KEY="secret",
    )
    @patch("qiniu.Auth.private_download_url", return_value="http://cdn.example.com/signed-image")
    def test_private_image_uses_signed_preview_url(self, mocked_sign):
        storage = QiniuStorageService()

        result = storage.access_url("docs/parsed/image.png")

        self.assertEqual(result, "http://cdn.example.com/signed-image")
        mocked_sign.assert_called_once()

    @override_settings(QINIU_ACCESS_KEY="access", QINIU_SECRET_KEY="secret", QINIU_BUCKET="bucket")
    @patch("qiniu.BucketManager.delete")
    def test_qiniu_delete_many_skips_empty_keys_and_deduplicates(self, mocked_delete):
        class Info:
            status_code = 200
            text_body = ""

        mocked_delete.return_value = ({}, Info())

        QiniuStorageService().delete_many(["docs/a.md", "", "docs/a.md", "docs/parsed/image.png", None])

        self.assertEqual(mocked_delete.call_count, 2)
        self.assertEqual(mocked_delete.call_args_list[0].args[1], "docs/a.md")
        self.assertEqual(mocked_delete.call_args_list[1].args[1], "docs/parsed/image.png")

    @override_settings(QINIU_ACCESS_KEY="access", QINIU_SECRET_KEY="secret", QINIU_BUCKET="bucket")
    @patch("qiniu.BucketManager.delete")
    def test_qiniu_delete_treats_missing_object_as_success(self, mocked_delete):
        class Info:
            status_code = 612
            text_body = "no such file"

        mocked_delete.return_value = ({}, Info())

        QiniuStorageService().delete("docs/missing.md")

        mocked_delete.assert_called_once()

    @patch.object(DocumentExtractionService, "_extract_with_docling", side_effect=RuntimeError("docling unavailable"))
    def test_docx_falls_back_and_preserves_headings_and_tables(self, _mocked_docling):
        import docx

        document = docx.Document()
        document.add_heading("登录功能", level=1)
        document.add_paragraph("用户可以使用账号密码登录。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "字段"
        table.cell(0, 1).text = "说明"
        table.cell(1, 0).text = "账号"
        table.cell(1, 1).text = "必填"
        with tempfile.NamedTemporaryFile(suffix=".docx") as temp:
            document.save(temp.name)
            result = DocumentExtractionService.extract_path(temp.name, "docx", "登录需求.docx")

        self.assertEqual(result["engine"], "legacy_docx")
        self.assertTrue(result["fallback_used"])
        self.assertEqual([block["type"] for block in result["blocks"]], ["heading", "paragraph", "table"])
        self.assertIn("| 字段 | 说明 |", result["plain_text"])
        self.assertEqual(result["blocks"][0]["source_locator"], "登录需求.docx#B1")

    def test_picture_blocks_are_not_added_to_plain_text(self):
        blocks = []
        DocumentExtractionService._append_block(blocks, "需求.pdf", "heading", "登录", 1, level=1)
        DocumentExtractionService._append_block(
            blocks, "需求.pdf", "picture", DocumentExtractionService.PICTURE_NOTE, 1
        )

        text = DocumentExtractionService._blocks_to_text(blocks)

        self.assertIn("登录", text)
        self.assertNotIn(DocumentExtractionService.PICTURE_NOTE, text)
        self.assertEqual(blocks[1]["source_locator"], "需求.pdf:P1#B2")

    def test_table_html_keeps_spans_and_removes_unsafe_markup(self):
        html = DocumentExtractionService._sanitize_table_html(
            '<table onclick="bad()"><tr><td rowspan="2"><script>alert(1)</script>内容</td></tr></table>'
        )

        self.assertIn('rowspan="2"', html)
        self.assertNotIn("script", html)
        self.assertNotIn("onclick", html)

    def test_table_cell_paragraphs_are_removed_but_following_body_is_kept(self):
        blocks = [
            {"type": "table", "rows": [["功能子项", "功能说明", "优先级"]]},
            {"type": "paragraph", "text": "功能子项"},
            {"type": "paragraph", "text": "功能说明"},
            {"type": "paragraph", "text": "优先级"},
            {"type": "paragraph", "text": "订单发布页面支持快速选择地址"},
        ]

        result = DocumentExtractionService._remove_duplicate_table_cells(blocks)

        self.assertEqual([block["type"] for block in result], ["table", "paragraph"])
        self.assertEqual(result[1]["text"], "订单发布页面支持快速选择地址")

    def test_structured_parser_uses_heading_hierarchy_and_keeps_orphans(self):
        blocks = [
            {"type": "heading", "text": "1 产品需求", "level": 1, "source_locator": "doc#B1"},
            {"type": "heading", "text": "项目背景", "level": 2, "source_locator": "doc#B2"},
            {"type": "paragraph", "text": "应被过滤", "source_locator": "doc#B3"},
            {"type": "heading", "text": "2.1 账号模块", "level": 2, "source_locator": "doc#B4"},
            {"type": "paragraph", "text": "模块说明", "source_locator": "doc#B5"},
            {"type": "heading", "text": "2.1.1 用户登录", "level": 3, "source_locator": "doc#B6"},
            {"type": "paragraph", "text": "用户使用密码登录", "source_locator": "doc#B7"},
            {"type": "heading", "text": "2.1.2 补充说明", "level": 4, "source_locator": "doc#B8"},
            {"type": "table", "text": "", "rows": [["错误", "提示"]], "source_locator": "doc#B9"},
        ]

        result = StructuredRequirementParser.organize(blocks, "默认标题")

        self.assertEqual(result["document_title"], "产品需求")
        self.assertEqual(len(result["requirements"]), 1)
        self.assertEqual(result["requirements"][0]["module"], "账号模块")
        self.assertEqual(result["requirements"][0]["title"], "用户登录")
        self.assertEqual(result["requirements"][0]["supplementary"], ["补充说明"])
        self.assertEqual(result["requirements"][0]["blocks"][-1]["block_type"], "table")
        self.assertEqual(result["orphan_blocks"][0]["text"], "模块说明")
        self.assertGreaterEqual(len(result["filtered"]), 2)


class RequirementCaseCenterTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="tester", password="password")
        self.client = APIClient()
        self.client.force_authenticate(self.user)
        self.project = ProjectConfig.objects.create(name="示例项目", code="demo", status="active")
        self.document = RequirementDocument.objects.create(
            project=self.project,
            title="需求文档",
            original_filename="requirements.md",
            document_type="md",
            file_size=128,
            qiniu_key="docs/requirements.md",
            uploaded_by=self.user,
        )
        self.item_a = RequirementItem.objects.create(
            project=self.project,
            document=self.document,
            requirement_no="REQ-001",
            title="登录",
            module="认证",
            description="用户可以登录系统",
            priority="medium",
            confirm_status="confirmed",
            confirmed_by=self.user,
            confirmed_at=timezone.now(),
        )
        self.item_b = RequirementItem.objects.create(
            project=self.project,
            document=self.document,
            requirement_no="REQ-002",
            title="退出",
            module="认证",
            description="用户可以退出系统",
            priority="medium",
            confirm_status="confirmed",
            confirmed_by=self.user,
            confirmed_at=timezone.now(),
        )

    def seed_ai_config(self):
        writer = LLMModelConfig.objects.create(
            name="writer",
            provider="deepseek",
            protocol="openai_compatible",
            usage="testcase_writer",
            model_name="writer-model",
            base_url="https://example.com",
            api_key="key",
            is_active=True,
            is_default=True,
        )
        reviewer = LLMModelConfig.objects.create(
            name="reviewer",
            provider="deepseek",
            protocol="openai_compatible",
            usage="testcase_reviewer",
            model_name="reviewer-model",
            base_url="https://example.com",
            api_key="key",
            is_active=True,
            is_default=True,
        )
        PromptConfig.objects.update_or_create(
            role_type="testcase_writer",
            defaults={"name": "writer prompt", "prompt_content": "生成用例", "llm_model": writer, "is_active": True},
        )
        PromptConfig.objects.update_or_create(
            role_type="testcase_reviewer",
            defaults={"name": "reviewer prompt", "prompt_content": "审核用例", "llm_model": reviewer, "is_active": True},
        )

    def seed_vision_config(self):
        model = LLMModelConfig.objects.create(
            name="vision",
            provider="chatgpt",
            protocol="openai_responses",
            usage="vision_analyzer",
            model_name="gpt-5-mini",
            base_url="https://api.openai.com",
            api_key="key",
            max_tokens=4096,
            temperature=0.1,
            top_p=1,
            is_active=True,
            is_default=True,
        )
        PromptConfig.objects.update_or_create(
            role_type="vision_analyzer",
            defaults={"name": "vision prompt", "prompt_content": "识别需求图片", "llm_model": model, "is_active": True},
        )
        return model

    def seed_gemini_vision_config(self):
        model = LLMModelConfig.objects.create(
            name="vision gemini",
            provider="gemini",
            protocol="gemini",
            usage="vision_analyzer",
            model_name="gemini-3.5-flash",
            base_url="https://generativelanguage.googleapis.com",
            api_key="key",
            max_tokens=4096,
            temperature=0.1,
            top_p=1,
            is_active=True,
            is_default=False,
        )
        PromptConfig.objects.update_or_create(
            role_type="vision_analyzer",
            defaults={"name": "vision prompt", "prompt_content": "识别需求图片", "llm_model": model, "is_active": True},
        )
        return model

    def create_integration_draft(self, item):
        return RequirementIntegrationDraft.objects.create(
            requirement_item=item,
            status="completed",
            title=item.title,
            module=item.module,
            description=item.description,
            acceptance_criteria=item.acceptance_criteria,
            supplementary_description=item.supplementary_description,
            source_summary="已整理",
            raw_context=f"需求编号: {item.requirement_no}",
            created_by=self.user,
        )

    def assign_formal_revisions(self, version, *items):
        module, _ = ProjectModule.objects.get_or_create(
            project=self.project,
            code="auth",
            defaults={"name": "认证"},
        )
        revisions = []
        for item in items:
            family = RequirementFamily.objects.create(
                project=self.project,
                family_no=f"FR-{item.id:03d}",
                title=item.title,
                created_by=self.user,
            )
            revision = RequirementRevision.objects.create(
                family=family,
                source_item=item,
                revision_no=1,
                change_type="initial",
                title=item.title,
                description=item.description,
                acceptance_criteria=item.acceptance_criteria,
                supplementary_description=item.supplementary_description,
                source_summary="已确认正式需求",
                source_content_hash=f"{item.id:064x}",
                confirmed_by=self.user,
            )
            family.modules.set([module])
            revision.modules.set([module])
            item.formal_modules.set([module])
            revisions.append(revision)
        version.requirement_items.add(*items)
        version.requirement_revisions.add(*revisions)
        if version.status == "draft":
            version.status = "published"
            version.published_by = self.user
            version.published_at = timezone.now()
            version.save(update_fields=["status", "published_by", "published_at", "updated_at"])
        return revisions

    @staticmethod
    def assign_task_requirements(task, version, *items):
        task.requirement_items.add(*items)
        task.requirement_revisions.set(
            version.requirement_revisions.filter(source_item__in=items)
        )

    def seed_integration_config(self):
        model = LLMModelConfig.objects.create(
            name="integrator",
            provider="deepseek",
            protocol="openai_compatible",
            usage="requirement_integrator",
            model_name="integrator-model",
            base_url="https://example.com",
            api_key="key",
            is_active=True,
            is_default=True,
        )
        return PromptConfig.objects.create(
            role_type="requirement_integrator",
            name="integrator prompt",
            prompt_content="整合需求",
            llm_model=model,
            is_active=True,
        )

    def test_created_version_is_draft(self):
        url = reverse("requirement-version-list")
        response = self.client.post(
            url,
            {
                "project": self.project.id,
                "version_no": "v1.0",
                "name": "首版",
            },
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        version = RequirementVersion.objects.get(pk=response.data["id"])
        self.assertEqual(version.status, "draft")
        self.assertIsNone(version.published_by)
        self.assertIsNone(version.published_at)

    def test_draft_version_can_bind_unbind_and_publish_requirements(self):
        version = RequirementVersion.objects.create(
            project=self.project,
            version_no="v1.0",
            name="首版",
            created_by=self.user,
        )
        revision = self.assign_formal_revisions(version, self.item_a)[0]
        version.status = "draft"
        version.published_by = None
        version.published_at = None
        version.save(update_fields=["status", "published_by", "published_at", "updated_at"])

        unbind_response = self.client.post(
            reverse("requirement-version-unbind-requirements", args=[version.id]),
            {"revision_ids": [revision.id]},
            format="json",
        )
        self.assertEqual(unbind_response.status_code, 200)
        self.assertFalse(version.requirement_revisions.exists())
        self.assertFalse(version.requirement_items.exists())

        empty_publish = self.client.post(reverse("requirement-version-publish", args=[version.id]), format="json")
        self.assertEqual(empty_publish.status_code, 409)

        bind_url = reverse("requirement-version-bind-requirements", args=[version.id])
        first_bind = self.client.post(bind_url, {"revision_ids": [revision.id]}, format="json")
        repeated_bind = self.client.post(bind_url, {"revision_ids": [revision.id]}, format="json")
        self.assertEqual(first_bind.status_code, 200)
        self.assertEqual(repeated_bind.status_code, 200)
        self.assertEqual(version.requirement_revisions.count(), 1)
        self.assertEqual(version.requirement_items.count(), 1)

        publish_response = self.client.post(reverse("requirement-version-publish", args=[version.id]), format="json")
        self.assertEqual(publish_response.status_code, 200)
        version.refresh_from_db()
        self.assertEqual(version.status, "published")
        self.assertEqual(version.published_by, self.user)
        self.assertIsNotNone(version.published_at)

    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_published_version_only_allows_appending_requirements(self, mocked_delay):
        version = RequirementVersion.objects.create(
            project=self.project,
            version_no="v1.0",
            name="首版",
            status="published",
            created_by=self.user,
            published_by=self.user,
            published_at=timezone.now(),
        )
        first_revision = self.assign_formal_revisions(version, self.item_a)[0]
        second_revision = self.assign_formal_revisions(version, self.item_b)[0]
        version.requirement_revisions.remove(second_revision)
        version.requirement_items.remove(self.item_b)

        bind_response = self.client.post(
            reverse("requirement-version-bind-requirements", args=[version.id]),
            {"revision_ids": [second_revision.id]},
            format="json",
        )
        unbind_response = self.client.post(
            reverse("requirement-version-unbind-requirements", args=[version.id]),
            {"revision_ids": [first_revision.id]},
            format="json",
        )

        self.assertEqual(bind_response.status_code, 200)
        self.assertEqual(unbind_response.status_code, 409)
        self.assertEqual(set(version.requirement_revisions.values_list("id", flat=True)), {first_revision.id, second_revision.id})
        self.assertEqual(set(version.requirement_items.values_list("id", flat=True)), {self.item_a.id, self.item_b.id})

        self.seed_ai_config()
        generation_response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_b.id]},
            format="json",
        )
        self.assertEqual(generation_response.status_code, 201)
        mocked_delay.assert_called_once()

    def test_version_binding_rejects_cross_project_and_archived_versions(self):
        other_project = ProjectConfig.objects.create(name="其他项目", code="other", status="active")
        other_document = RequirementDocument.objects.create(
            project=other_project,
            title="其他需求",
            original_filename="other.md",
            qiniu_key="docs/other.md",
            uploaded_by=self.user,
        )
        other_item = RequirementItem.objects.create(
            project=other_project,
            document=other_document,
            requirement_no="OTHER-001",
            title="其他需求",
            module="其他",
            description="其他项目需求",
            confirm_status="confirmed",
            confirmed_by=self.user,
            confirmed_at=timezone.now(),
        )
        other_module = ProjectModule.objects.create(project=other_project, code="other", name="其他")
        other_family = RequirementFamily.objects.create(
            project=other_project,
            family_no="OTHER-FR-001",
            title="其他需求",
            created_by=self.user,
        )
        other_revision = RequirementRevision.objects.create(
            family=other_family,
            source_item=other_item,
            revision_no=1,
            change_type="initial",
            title="其他需求",
            description="其他项目需求",
            source_content_hash="b" * 64,
            confirmed_by=self.user,
        )
        other_family.modules.set([other_module])
        other_revision.modules.set([other_module])
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        bind_url = reverse("requirement-version-bind-requirements", args=[version.id])

        cross_project_response = self.client.post(bind_url, {"revision_ids": [other_revision.id]}, format="json")
        version.status = "archived"
        version.save(update_fields=["status", "updated_at"])
        archived_response = self.client.post(bind_url, {"revision_ids": [other_revision.id]}, format="json")

        self.assertEqual(cross_project_response.status_code, 400)
        self.assertEqual(archived_response.status_code, 409)

    def test_requirement_items_can_filter_by_version_with_pagination(self):
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        other_version = RequirementVersion.objects.create(project=self.project, version_no="v2.0", name="二版", created_by=self.user)
        version.requirement_items.set([self.item_a])
        other_version.requirement_items.set([self.item_b])

        response = self.client.get(
            reverse("requirement-item-list"),
            {"project": self.project.id, "version": version.id, "page": 1, "page_size": 10},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["count"], 1)
        self.assertEqual(response.data["results"][0]["id"], self.item_a.id)

    def test_requirement_item_list_hides_unconfirmed_by_default(self):
        self.item_a.confirm_status = "pending"
        self.item_a.confirmed_by = None
        self.item_a.confirmed_at = None
        self.item_a.save(update_fields=["confirm_status", "confirmed_by", "confirmed_at", "updated_at"])

        response = self.client.get(reverse("requirement-item-list"), {"project": self.project.id, "page": 1, "page_size": 10})

        self.assertEqual(response.status_code, 200)
        ids = [item["id"] for item in response.data["results"]]
        self.assertNotIn(self.item_a.id, ids)
        self.assertIn(self.item_b.id, ids)

    def test_requirement_item_list_can_include_unconfirmed_for_diagnostics(self):
        self.item_a.confirm_status = "pending"
        self.item_a.confirmed_by = None
        self.item_a.confirmed_at = None
        self.item_a.save(update_fields=["confirm_status", "confirmed_by", "confirmed_at", "updated_at"])

        response = self.client.get(
            reverse("requirement-item-list"),
            {"project": self.project.id, "include_unconfirmed": "true", "page": 1, "page_size": 10},
        )

        self.assertEqual(response.status_code, 200)
        ids = [item["id"] for item in response.data["results"]]
        self.assertIn(self.item_a.id, ids)

    @patch("apps.requirements.integration.RequirementReviewService._search", return_value=[])
    @patch("apps.requirements.services.LLMChatService.chat")
    def test_requirement_item_can_generate_integration_draft(self, mocked_chat, _mocked_search):
        self.seed_integration_config()
        ProjectModule.objects.create(project=self.project, code="auth", name="认证")
        mocked_chat.return_value = json.dumps({
            "title": "登录整合",
            "module_paths": ["认证"],
            "description": "用户使用账号密码登录系统",
            "acceptance_criteria": "登录成功后进入首页",
            "supplementary_description": "包含异常提示",
            "source_summary": "来自需求描述和来源内容块",
            "relationship_mode": "new",
            "change_type": "initial",
            "selected_revision_id": "",
            "conflicts": [],
            "open_questions": [],
            "evidence_refs": [],
        }, ensure_ascii=False)

        response = self.client.post(reverse("requirement-item-integrate", args=[self.item_a.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["status"], "completed")
        draft = RequirementIntegrationDraft.objects.get(requirement_item=self.item_a)
        self.assertEqual(draft.title, "登录整合")
        self.assertEqual(draft.model_name, "integrator-model")
        self.assertIn("需求编号: REQ-001", draft.raw_context)

    def test_requirement_item_integration_patch_updates_editable_fields(self):
        draft = RequirementIntegrationDraft.objects.create(
            requirement_item=self.item_a,
            status="completed",
            title="旧标题",
            module="认证",
            description="旧描述",
            created_by=self.user,
        )

        response = self.client.patch(
            reverse("requirement-item-integration", args=[self.item_a.id]),
            {"title": "手动标题", "description": "手动整理后的描述", "status": "failed"},
            format="json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["formal_module_ids"], [])
        draft.refresh_from_db()
        self.assertEqual(draft.title, "手动标题")
        self.assertEqual(draft.description, "手动整理后的描述")
        self.assertEqual(draft.status, "completed")
        self.assertEqual(draft.updated_by, self.user)

    @patch("apps.requirements.integration.RequirementReviewService._search", return_value=[])
    def test_requirement_item_integrate_requires_integrator_role(self, _mocked_search):
        response = self.client.post(reverse("requirement-item-integrate", args=[self.item_a.id]))

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.data["error"]["code"], "MODEL_CONFIGURATION_MISSING")
        self.assertIn("缺少启用的系统角色: 需求整合专家", response.data["detail"])

    def test_confirm_requirement_items_requires_completed_integration_review(self):
        self.item_a.confirm_status = "pending"
        self.item_a.confirmed_by = None
        self.item_a.confirmed_at = None
        self.item_a.save(update_fields=["confirm_status", "confirmed_by", "confirmed_at", "updated_at"])

        response = self.client.post(reverse("requirement-item-confirm"), {"ids": [self.item_a.id]}, format="json")

        self.assertEqual(response.status_code, 409)
        self.assertIn("请先完成需求整合", response.data["detail"])
        self.item_a.refresh_from_db()
        self.assertEqual(self.item_a.confirm_status, "pending")
        self.assertIsNone(self.item_a.confirmed_by)
        self.assertIsNone(self.item_a.confirmed_at)

    def test_confirm_requirement_items_rejects_archived_items(self):
        self.item_a.is_current = False
        self.item_a.is_archived = True
        self.item_a.save(update_fields=["is_current", "is_archived", "updated_at"])

        response = self.client.post(reverse("requirement-item-confirm"), {"ids": [self.item_a.id]}, format="json")

        self.assertEqual(response.status_code, 400)
        self.assertIn("只能确认当前有效且未归档的需求", response.data["detail"])

    @patch("apps.requirements.views.QiniuStorageService.delete_many")
    def test_document_delete_removes_qiniu_original_before_database_record(self, mocked_delete_many):
        response = self.client.delete(reverse("requirement-document-detail", args=[self.document.id]))

        self.assertEqual(response.status_code, 204)
        mocked_delete_many.assert_called_once_with(["docs/requirements.md"])
        self.assertFalse(RequirementDocument.objects.filter(pk=self.document.id).exists())

    @patch("apps.requirements.views.QiniuStorageService.delete_many")
    def test_document_delete_removes_parsed_image_keys(self, mocked_delete_many):
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", created_by=self.user)
        RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_a, block_type="image", order=1, image_key="docs/parsed/1.png")
        RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_a, block_type="image", order=2, image_key="")
        RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_b, block_type="image", order=3, image_key="docs/parsed/1.png")
        RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_b, block_type="image", order=4, image_key="docs/parsed/2.png")

        response = self.client.delete(reverse("requirement-document-detail", args=[self.document.id]))

        self.assertEqual(response.status_code, 204)
        keys = mocked_delete_many.call_args.args[0]
        self.assertEqual(keys[0], "docs/requirements.md")
        self.assertNotIn("", keys)
        self.assertEqual(set(keys), {"docs/requirements.md", "docs/parsed/1.png", "docs/parsed/2.png"})
        self.assertFalse(RequirementDocument.objects.filter(pk=self.document.id).exists())

    @patch("apps.requirements.views.QiniuStorageService.delete_many", side_effect=RuntimeError("docs/requirements.md: forbidden"))
    def test_document_delete_keeps_database_record_when_qiniu_delete_fails(self, mocked_delete_many):
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", created_by=self.user)
        block = RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_a, block_type="image", order=1, image_key="docs/parsed/1.png")

        response = self.client.delete(reverse("requirement-document-detail", args=[self.document.id]))

        self.assertEqual(response.status_code, 400)
        self.assertIn("七牛文件删除失败", response.data["detail"])
        mocked_delete_many.assert_called_once()
        self.assertTrue(RequirementDocument.objects.filter(pk=self.document.id).exists())
        self.assertTrue(RequirementParseRun.objects.filter(pk=run.id).exists())
        self.assertTrue(RequirementContentBlock.objects.filter(pk=block.id).exists())

    def test_document_list_does_not_return_extracted_blocks(self):
        self.document.extraction_engine = "docling"
        self.document.extracted_blocks = [
            {"type": "heading", "text": "登录", "block_index": 1},
            {"type": "table", "text": "|字段|", "block_index": 2},
            {"type": "picture", "text": "图片内容暂未识别", "block_index": 3},
        ]
        self.document.save(update_fields=["extraction_engine", "extracted_blocks"])

        response = self.client.get(reverse("requirement-document-list"), {"project": self.project.id})

        self.assertEqual(response.status_code, 200)
        payload = response.data["results"][0]
        self.assertNotIn("extracted_blocks", payload)
        self.assertIsNone(payload["current_run"])

    def test_document_content_returns_current_requirements_and_orphans(self):
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", is_current=True, created_by=self.user)
        self.item_a.parse_run = run
        self.item_a.save(update_fields=["parse_run"])
        RequirementContentBlock.objects.create(parse_run=run, requirement=self.item_a, block_type="text", order=1, text="用户可以登录")
        RequirementContentBlock.objects.create(parse_run=run, block_type="table", order=2, table_data={"rows": [["字段"]]})

        response = self.client.get(reverse("requirement-document-content", args=[self.document.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["requirements"][0]["content_blocks"][0]["text"], "用户可以登录")
        self.assertEqual(response.data["orphan_blocks"][0]["block_type"], "table")

    @patch("apps.requirements.services.urlopen")
    def test_vision_service_sends_input_image_to_responses_api(self, mocked_urlopen):
        mocked_urlopen.return_value = FakeUrlopenResponse({
            "output": [{"content": [{"text": '{"image_type":"ui","ui_elements":[],"business_rules":[],"flows":[],"states":[],"test_points":["登录按钮"],"uncertainties":[]}'}]}],
        })
        config = self.seed_vision_config()

        output_text, _raw = LLMResponsesVisionService(config).analyze_image(
            "https://cdn.example.com/image.png",
            "识别图片",
        )

        self.assertIn("登录按钮", output_text)
        request = mocked_urlopen.call_args.args[0]
        self.assertEqual(request.full_url, "https://api.openai.com/v1/responses")
        body = json.loads(request.data.decode("utf-8"))
        content = body["input"][0]["content"]
        self.assertEqual(content[0]["type"], "input_text")
        self.assertEqual(content[1]["type"], "input_image")
        self.assertEqual(content[1]["image_url"], "https://cdn.example.com/image.png")

    @patch("apps.requirements.services.urlopen")
    def test_gemini_vision_service_sends_inline_image_data(self, mocked_urlopen):
        mocked_urlopen.side_effect = [
            FakeUrlopenResponse(b"image-bytes", content_type="image/png"),
            FakeUrlopenResponse({
                "candidates": [
                    {
                        "content": {
                            "parts": [
                                {"text": '{"image_type":"ui","ui_elements":[],"business_rules":[],"flows":[],"states":[],"test_points":["登录"],"uncertainties":[]}'}
                            ]
                        }
                    }
                ]
            }),
        ]
        config = self.seed_gemini_vision_config()

        output_text, _raw = LLMGeminiVisionService(config).analyze_image(
            "https://cdn.example.com/image.png",
            "识别图片",
        )

        self.assertIn("登录", output_text)
        request = mocked_urlopen.call_args_list[1].args[0]
        self.assertEqual(
            request.full_url,
            "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key=key",
        )
        body = json.loads(request.data.decode("utf-8"))
        parts = body["contents"][0]["parts"]
        self.assertEqual(parts[0]["text"], "识别图片")
        self.assertEqual(parts[1]["inline_data"]["mime_type"], "image/png")
        self.assertEqual(parts[1]["inline_data"]["data"], "aW1hZ2UtYnl0ZXM=")

    def test_active_non_default_vision_model_can_be_selected(self):
        config = self.seed_gemini_vision_config()

        selected = RequirementImageAnalysisService.get_vision_config()

        self.assertEqual(selected, config)

    @patch("apps.requirements.services.LLMResponsesVisionService.analyze_image")
    @patch("apps.requirements.services.QiniuStorageService.access_url", return_value="https://cdn.example.com/image.png")
    def test_image_analysis_is_cached_for_same_block(self, _mocked_access_url, mocked_analyze):
        self.seed_vision_config()
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", is_current=True, created_by=self.user)
        block = RequirementContentBlock.objects.create(
            parse_run=run,
            requirement=self.item_a,
            block_type="image",
            order=1,
            image_key="docs/parsed/image.png",
            source_locator="requirements.md#B1",
        )
        mocked_analyze.return_value = (
            '{"image_type":"ui","ui_elements":["登录按钮"],"business_rules":[],"flows":[],"states":[],"test_points":["点击登录"],"uncertainties":[]}',
            {"id": "resp_1", "model": "gpt-5-mini"},
        )

        RequirementImageAnalysisService.ensure_for_requirement(self.item_a)
        RequirementImageAnalysisService.ensure_for_requirement(self.item_a)

        self.assertEqual(mocked_analyze.call_count, 1)
        analysis = RequirementImageAnalysis.objects.get(content_block=block)
        self.assertEqual(analysis.status, "completed")
        self.assertEqual(analysis.summary["test_points"], ["点击登录"])

    def test_requirement_context_includes_table_and_image_summary(self):
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", is_current=True, created_by=self.user)
        table_block = RequirementContentBlock.objects.create(
            parse_run=run,
            requirement=self.item_a,
            block_type="table",
            order=1,
            table_data={"rows": [["字段", "规则"], ["手机号", "必填"]]},
            source_locator="requirements.md#B2",
        )
        image_block = RequirementContentBlock.objects.create(
            parse_run=run,
            requirement=self.item_a,
            block_type="image",
            order=2,
            image_key="docs/parsed/image.png",
            source_locator="requirements.md#B3",
        )
        RequirementImageAnalysis.objects.create(
            content_block=image_block,
            status="completed",
            model_name="gpt-5-mini",
            summary={
                "image_type": "ui",
                "ui_elements": ["登录按钮"],
                "business_rules": [],
                "flows": [],
                "states": [],
                "test_points": ["验证点击登录按钮"],
                "uncertainties": [],
            },
        )

        context = RequirementContextBuilder.build(self.item_a)

        self.assertIn("| 字段 | 规则 |", context)
        self.assertIn("验证点击登录按钮", context)
        self.assertIn(table_block.source_locator, context)

    def test_parse_test_cases_recovers_complete_items_from_truncated_json_array(self):
        content = (
            '[{"case_no":"TC-001","title":"验证新用户通过手机号验证码注册并自动登录成功",'
            '"preconditions":"测试手机号未在系统注册","steps":"1. 打开小程序登录页面\\n2. 输入手机号",'
            '"expected_result":"登录成功","priority":"high","test_type":"functional"},'
            '{"case_no":"TC-002","title":"验证老用户通过手机号验证码快速登录成功",'
            '"preconditions":"测试手机号已在系统注册","steps":"1. 打开小程序登录页面\\n2. 输入已注册手机号'
        )

        cases = TestCaseGenerationService.parse_test_cases(content)

        self.assertEqual(len(cases), 1)
        self.assertEqual(cases[0]["case_no"], "TC-001")
        self.assertEqual(cases[0]["priority"], "high")

    @patch("apps.requirements.services.LLMChatService.chat")
    def test_generate_all_for_requirement_polls_until_less_than_batch_size(self, mocked_chat):
        self.seed_ai_config()

        def render_cases(start, count):
            return [
                {
                    "case_no": f"TC-{index:03d}",
                    "title": f"场景 {index}",
                    "preconditions": "无",
                    "steps": f"执行步骤 {index}",
                    "expected_result": f"预期结果 {index}",
                    "priority": "medium",
                    "test_type": "functional",
                }
                for index in range(start, start + count)
            ]

        mocked_chat.side_effect = [
            json.dumps(render_cases(1, TestCaseGenerationService.CASE_BATCH_SIZE), ensure_ascii=False),
            json.dumps(render_cases(7, 2), ensure_ascii=False),
        ]

        cases, _raw, _writer_model, _writer_role, rounds = TestCaseGenerationService.generate_all_for_requirement(self.item_a)

        self.assertEqual(rounds, 2)
        self.assertEqual(len(cases), 8)
        self.assertEqual(cases[-1]["case_no"], "TC-008")
        self.assertIn("已生成用例摘要", mocked_chat.call_args_list[1].args[1])

    @patch("apps.requirements.services.LLMChatService.chat")
    def test_generate_all_for_requirement_stops_when_model_returns_no_new_cases(self, mocked_chat):
        self.seed_ai_config()
        first_batch = [
            {
                "case_no": f"TC-{index:03d}",
                "title": f"场景 {index}",
                "preconditions": "无",
                "steps": f"执行步骤 {index}",
                "expected_result": f"预期结果 {index}",
                "priority": "medium",
                "test_type": "functional",
            }
            for index in range(1, TestCaseGenerationService.CASE_BATCH_SIZE + 1)
        ]
        mocked_chat.side_effect = [
            json.dumps(first_batch, ensure_ascii=False),
            "[]",
        ]

        cases, _raw, _writer_model, _writer_role, rounds = TestCaseGenerationService.generate_all_for_requirement(self.item_a)

        self.assertEqual(rounds, 2)
        self.assertEqual(len(cases), TestCaseGenerationService.CASE_BATCH_SIZE)

    @patch.object(DocumentExtractionService, "extract_bytes")
    @patch("apps.requirements.views.QiniuStorageService.download")
    def test_reparse_downloads_original_and_rebuilds_blocks_and_items(self, mocked_download, mocked_extract):
        mocked_download.return_value = b"new document bytes"
        mocked_extract.return_value = {
            "plain_text": "# 产品需求\n\n## 认证\n\n### 新登录需求\n\n用户必须使用验证码登录。",
            "blocks": [
                {"type": "heading", "text": "产品需求", "page": 2, "block_index": 1,
                 "level": 1, "source_locator": "requirements.md:P2#B1"},
                {"type": "heading", "text": "认证", "page": 2, "block_index": 2,
                 "level": 2, "source_locator": "requirements.md:P2#B2"},
                {"type": "heading", "text": "新登录需求", "page": 2, "block_index": 3,
                 "level": 3, "source_locator": "requirements.md:P2#B3"},
                {"type": "paragraph", "text": "用户必须使用验证码登录。", "page": 2,
                 "block_index": 4, "source_locator": "requirements.md:P2#B4"},
            ],
            "engine": "docling",
            "fallback_used": False,
        }

        response = self.client.post(reverse("requirement-document-parse", args=[self.document.id]))

        self.assertEqual(response.status_code, 200)
        mocked_download.assert_called_once_with(self.document.qiniu_key, self.document.qiniu_url)
        mocked_extract.assert_called_once_with(
            b"new document bytes", self.document.document_type, self.document.original_filename
        )
        self.document.refresh_from_db()
        self.assertEqual(self.document.extraction_engine, "docling")
        self.assertEqual(self.document.extracted_blocks[0]["page"], 2)
        self.assertEqual(self.document.items.filter(is_current=True).count(), 1)
        self.assertTrue(self.item_a.__class__.objects.get(pk=self.item_a.pk).is_archived)
        rebuilt = self.document.items.get(is_current=True)
        self.assertEqual(rebuilt.title, "新登录需求")
        self.assertEqual(rebuilt.confirm_status, "pending")
        self.assertIsNone(rebuilt.confirmed_by)
        self.assertIsNone(rebuilt.confirmed_at)

    def test_published_version_can_be_archived_but_not_edited(self):
        version = RequirementVersion.objects.create(
            project=self.project,
            version_no="v1.0",
            name="首版",
            status="published",
            created_by=self.user,
            published_by=self.user,
            published_at=timezone.now(),
        )

        edit_response = self.client.patch(
            reverse("requirement-version-detail", args=[version.id]),
            {"name": "修改后的名称"},
            format="json",
        )
        archive_response = self.client.post(
            reverse("requirement-version-archive", args=[version.id]),
            format="json",
        )

        self.assertEqual(edit_response.status_code, 405)
        self.assertEqual(archive_response.status_code, 200)
        version.refresh_from_db()
        self.assertEqual(version.status, "archived")

    def test_generation_request_requires_default_ai_configuration(self):
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("configuration", response.data)

    def test_generation_request_rejects_unconfirmed_requirement_items(self):
        self.seed_ai_config()
        self.item_a.confirm_status = "pending"
        self.item_a.confirmed_by = None
        self.item_a.confirmed_at = None
        self.item_a.save(update_fields=["confirm_status", "confirmed_by", "confirmed_at", "updated_at"])
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("尚未确认", str(response.data))

    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_generation_request_creates_background_task(self, mocked_delay):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        task = TestCaseGenerationTask.objects.get(pk=response.data["id"])
        self.assertEqual(task.total_count, 1)
        mocked_delay.assert_called_once_with(task.id)

    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_failed_generation_retry_creates_new_linked_task(self, mocked_delay):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v-retry", name="重试版本", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        source = TestCaseGenerationTask.objects.create(
            task_no="TCG-FAILED-RETRY",
            project=self.project,
            version=version,
            total_count=1,
            failed_count=1,
            status="failed",
            generation_log=[{
                "requirement_item": self.item_a.id,
                "requirement_no": self.item_a.requirement_no,
                "status": "failed",
                "stage": "模型生成",
                "message": "模型服务当前繁忙",
            }],
            created_by=self.user,
        )
        self.assign_task_requirements(source, version, self.item_a)

        response = self.client.post(reverse("requirement-generation-task-retry", args=[source.id]), format="json")

        self.assertEqual(response.status_code, 201)
        retry_task = TestCaseGenerationTask.objects.get(pk=response.data["id"])
        self.assertEqual(retry_task.retry_of_id, source.id)
        self.assertEqual(list(retry_task.requirement_items.values_list("id", flat=True)), [self.item_a.id])
        mocked_delay.assert_called_once_with(retry_task.id)

    @patch("apps.requirements.tasks.run_requirement_integration_batch.delay")
    def test_partial_integration_retry_only_copies_failed_items(self, mocked_delay):
        source = RequirementIntegrationBatch.objects.create(
            project=self.project,
            document=self.document,
            status="partial_success",
            total_count=2,
            success_count=1,
            failed_count=1,
            created_by=self.user,
        )
        source.requirement_items.set([self.item_a, self.item_b])
        RequirementIntegrationRun.objects.create(
            batch=source,
            requirement_item=self.item_b,
            status="failed",
            source_content_hash="b" * 64,
            created_by=self.user,
        )

        response = self.client.post(reverse("requirement-integration-batch-retry", args=[source.id]), format="json")

        self.assertEqual(response.status_code, 201)
        retry_batch = RequirementIntegrationBatch.objects.get(pk=response.data["id"])
        self.assertEqual(retry_batch.retry_of_id, source.id)
        self.assertEqual(list(retry_batch.requirement_items.values_list("id", flat=True)), [self.item_b.id])
        mocked_delay.assert_called_once_with(retry_batch.id)

    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_generation_request_reuses_active_task_for_same_requirement_items(self, mocked_delay):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-EXISTING",
            project=self.project,
            version=version,
            total_count=1,
            status="running",
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["id"], task.id)
        self.assertEqual(TestCaseGenerationTask.objects.count(), 1)
        mocked_delay.assert_not_called()

    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_generation_request_marks_stale_pending_task_failed_before_creating_new_task(self, mocked_delay):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        stale_task = TestCaseGenerationTask.objects.create(
            task_no="TCG-STALE",
            project=self.project,
            version=version,
            total_count=1,
            status="pending",
            created_by=self.user,
        )
        TestCaseGenerationTask.objects.filter(pk=stale_task.pk).update(created_at=timezone.now() - timedelta(minutes=3))
        self.assign_task_requirements(stale_task, version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        stale_task.refresh_from_db()
        self.assertEqual(stale_task.status, "failed")
        self.assertEqual(stale_task.error_message, "后台任务队列不可用")
        self.assertEqual(stale_task.error_info["code"], "QUEUE_UNAVAILABLE")
        self.assertEqual(TestCaseGenerationTask.objects.count(), 2)

    @patch("apps.requirements.views.threading.Thread")
    @patch("apps.requirements.tasks.run_testcase_generation_task.delay")
    def test_generation_request_falls_back_to_local_thread_when_broker_is_unavailable(self, mocked_delay, mocked_thread):
        self.seed_ai_config()
        mocked_delay.side_effect = OperationalError("Connection refused")
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)

        response = self.client.post(
            reverse("requirement-generation-task-generate"),
            {"project": self.project.id, "version": version.id, "requirement_items": [self.item_a.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        task = TestCaseGenerationTask.objects.get(pk=response.data["id"])
        mocked_delay.assert_called_once_with(task.id)
        mocked_thread.assert_called_once()
        self.assertTrue(mocked_thread.call_args.kwargs["daemon"])
        mocked_thread.return_value.start.assert_called_once()

    @patch("apps.requirements.tasks.TestCaseGenerationService.review_cases")
    @patch("apps.requirements.tasks.TestCaseGenerationService.generate_all_for_requirement")
    def test_task_continues_after_single_requirement_failure_and_retries_review_failure(self, mocked_generate, mocked_review):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a, self.item_b)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-001",
            project=self.project,
            version=version,
            total_count=2,
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a, self.item_b)
        self.create_integration_draft(self.item_a)
        self.create_integration_draft(self.item_b)
        writer = LLMModelConfig.objects.get(usage="testcase_writer")
        reviewer = LLMModelConfig.objects.get(usage="testcase_reviewer")
        writer_role = PromptConfig.objects.get(role_type="testcase_writer")
        reviewer_role = PromptConfig.objects.get(role_type="testcase_reviewer")
        case = {
            "case_no": "TC-001",
            "title": "登录成功",
            "preconditions": "存在用户",
            "steps": "输入账号密码",
            "expected_result": "登录成功",
            "priority": "medium",
            "test_type": "functional",
            "raw": {},
        }
        mocked_generate.side_effect = [
            ([case], "raw", writer, writer_role, 1),
            ([case], "raw", writer, writer_role, 1),
            Exception("boom"),
        ]
        mocked_review.side_effect = [
            (False, "缺少异常场景", reviewer, reviewer_role),
            (True, "通过", reviewer, reviewer_role),
        ]

        run_testcase_generation_task(task.id)

        task.refresh_from_db()
        self.assertEqual(task.status, "partial_success")
        self.assertEqual(task.success_count, 1)
        self.assertEqual(task.failed_count, 1)
        self.assertEqual(LibraryTestCase.objects.count(), 1)
        self.assertEqual(mocked_generate.call_count, 3)
        self.assertEqual(task.generation_log[0]["writer_role"], "writer prompt")
        self.assertEqual(task.generation_log[0]["reviewer_role"], "reviewer prompt")
        self.assertEqual(task.generation_log[0]["stage"], "完成")
        self.assertEqual(task.generation_log[1]["stage"], "失败")
        self.assertEqual(task.generation_log[1]["message"], "系统内部异常")
        self.assertEqual(task.generation_log[1]["error"]["code"], "INTERNAL_ERROR")
        self.assertNotIn("boom", task.generation_log[1]["message"])

    @patch("apps.requirements.tasks.TestCaseGenerationService.review_cases")
    @patch("apps.requirements.tasks.TestCaseGenerationService.generate_all_for_requirement")
    def test_task_keeps_first_cases_when_retry_generation_is_unparseable(self, mocked_generate, mocked_review):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-RETRY-UNPARSEABLE",
            project=self.project,
            version=version,
            total_count=1,
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a)
        self.create_integration_draft(self.item_a)
        writer = LLMModelConfig.objects.get(usage="testcase_writer")
        reviewer = LLMModelConfig.objects.get(usage="testcase_reviewer")
        writer_role = PromptConfig.objects.get(role_type="testcase_writer")
        reviewer_role = PromptConfig.objects.get(role_type="testcase_reviewer")
        case = {
            "case_no": "TC-001",
            "title": "登录成功",
            "preconditions": "存在用户",
            "steps": "输入账号密码",
            "expected_result": "登录成功",
            "priority": "medium",
            "test_type": "functional",
            "raw": {},
        }
        mocked_generate.side_effect = [
            ([case], "raw", writer, writer_role, 1),
            TestCaseGenerationError("模型未返回可解析的测试用例"),
        ]
        mocked_review.return_value = (False, "缺少异常场景", reviewer, reviewer_role)

        run_testcase_generation_task(task.id)

        task.refresh_from_db()
        self.assertEqual(task.status, "completed")
        self.assertEqual(task.success_count, 1)
        self.assertEqual(task.failed_count, 0)
        self.assertEqual(LibraryTestCase.objects.count(), 1)
        self.assertTrue(task.generation_log[0]["retry_failed"])
        self.assertIn("重新生成失败", LibraryTestCase.objects.get().review_feedback)

    @patch("apps.requirements.tasks.TestCaseGenerationService.review_cases")
    @patch("apps.requirements.tasks.TestCaseGenerationService.generate_all_for_requirement")
    def test_task_uses_formal_revision_without_integration_draft(self, mocked_generate, mocked_review):
        self.seed_ai_config()
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-AUTO-INTEGRATE",
            project=self.project,
            version=version,
            total_count=1,
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a)
        draft = self.create_integration_draft(self.item_a)
        writer = LLMModelConfig.objects.get(usage="testcase_writer")
        reviewer = LLMModelConfig.objects.get(usage="testcase_reviewer")
        writer_role = PromptConfig.objects.get(role_type="testcase_writer")
        reviewer_role = PromptConfig.objects.get(role_type="testcase_reviewer")
        case = {
            "case_no": "TC-001",
            "title": "登录成功",
            "preconditions": "存在用户",
            "steps": "输入账号密码",
            "expected_result": "登录成功",
            "priority": "medium",
            "test_type": "functional",
            "raw": {},
        }
        mocked_generate.return_value = ([case], "raw", writer, writer_role, 1)
        mocked_review.return_value = (True, "通过", reviewer, reviewer_role)

        RequirementIntegrationDraft.objects.filter(pk=draft.pk).delete()
        run_testcase_generation_task(task.id)

        task.refresh_from_db()
        self.assertEqual(task.status, "completed")
        self.assertIn("来源摘要:\n已确认正式需求", mocked_generate.call_args.kwargs["requirement_context"])

    @patch("apps.requirements.tasks.TestCaseGenerationService.review_cases")
    @patch("apps.requirements.tasks.TestCaseGenerationService.generate_all_for_requirement")
    def test_task_uses_immutable_formal_revision_instead_of_editable_draft(self, mocked_generate, mocked_review):
        self.seed_ai_config()
        draft = self.create_integration_draft(self.item_a)
        draft.description = "手动整合后的需求"
        draft.save(update_fields=["description", "updated_at"])
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-REUSE-INTEGRATE",
            project=self.project,
            version=version,
            total_count=1,
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a)
        writer = LLMModelConfig.objects.get(usage="testcase_writer")
        reviewer = LLMModelConfig.objects.get(usage="testcase_reviewer")
        writer_role = PromptConfig.objects.get(role_type="testcase_writer")
        reviewer_role = PromptConfig.objects.get(role_type="testcase_reviewer")
        case = {
            "case_no": "TC-001",
            "title": "登录成功",
            "preconditions": "存在用户",
            "steps": "输入账号密码",
            "expected_result": "登录成功",
            "priority": "medium",
            "test_type": "functional",
            "raw": {},
        }
        mocked_generate.return_value = ([case], "raw", writer, writer_role, 1)
        mocked_review.return_value = (True, "通过", reviewer, reviewer_role)

        run_testcase_generation_task(task.id)

        requirement_context = mocked_generate.call_args.kwargs["requirement_context"]
        self.assertIn("用户可以登录系统", requirement_context)
        self.assertNotIn("手动整合后的需求", requirement_context)

    def test_task_fails_image_requirement_without_active_vision_role(self):
        self.seed_ai_config()
        run = RequirementParseRun.objects.create(document=self.document, run_no=1, status="completed", is_current=True, created_by=self.user)
        self.item_a.parse_run = run
        self.item_a.save(update_fields=["parse_run"])
        RequirementContentBlock.objects.create(
            parse_run=run,
            requirement=self.item_a,
            block_type="image",
            order=1,
            image_key="docs/parsed/image.png",
        )
        version = RequirementVersion.objects.create(project=self.project, version_no="v1.0", name="首版", created_by=self.user)
        self.assign_formal_revisions(version, self.item_a)
        task = TestCaseGenerationTask.objects.create(
            task_no="TCG-VISION-MISSING",
            project=self.project,
            version=version,
            total_count=1,
            created_by=self.user,
        )
        self.assign_task_requirements(task, version, self.item_a)

        run_testcase_generation_task(task.id)

        task.refresh_from_db()
        self.assertEqual(task.status, "failed")
        self.assertEqual(task.generation_log[0]["stage"], "失败")
        self.assertIn("缺少启用的系统角色: 图片理解专家", task.generation_log[0]["message"])


class RequirementIntegrationWorkflowTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="workflow-tester", password="password")
        self.client = APIClient()
        self.client.force_authenticate(self.user)
        self.project = ProjectConfig.objects.create(name="流程项目", code="workflow", status="active")
        self.document = RequirementDocument.objects.create(
            project=self.project,
            title="流程文档",
            original_filename="workflow.md",
            document_type="md",
            file_size=128,
            qiniu_key="docs/workflow.md",
            uploaded_by=self.user,
        )
        self.current_item = RequirementItem.objects.create(
            project=self.project,
            document=self.document,
            requirement_no="REQ-001",
            title="当前需求",
            module="认证",
            description="当前待整合需求",
        )
        self.sibling_item = RequirementItem.objects.create(
            project=self.project,
            document=self.document,
            requirement_no="REQ-002",
            title="同文档需求",
            module="认证",
            description="同一文档中的关联需求",
        )

    def test_integration_context_contains_other_current_document_requirements(self):
        context = RequirementReviewService._current_document_context(self.current_item)

        self.assertEqual([item["id"] for item in context], [self.sibling_item.id])
        self.assertEqual(context[0]["description"], "同一文档中的关联需求")

    @patch("apps.requirements.views.RequirementReviewService.integrate", side_effect=RuntimeError("integration reached"))
    def test_pending_parsed_requirement_can_reach_integrate_action(self, mocked_integrate):
        response = self.client.post(reverse("requirement-item-integrate", args=[self.current_item.id]), format="json")

        self.assertEqual(response.status_code, 500)
        self.assertEqual(response.data["error"]["code"], "INTERNAL_ERROR")
        self.assertNotIn("integration reached", response.data["detail"])
        mocked_integrate.assert_called_once_with(self.current_item, self.user)

    def test_new_requirement_accepts_empty_selected_revision_id(self):
        module = ProjectModule.objects.create(project=self.project, code="auth", name="认证")
        run = RequirementIntegrationRun.objects.create(
            requirement_item=self.current_item,
            status="running",
            source_content_hash=RequirementReviewService.source_hash(self.current_item),
            created_by=self.user,
        )
        role = SimpleNamespace(name="需求整合专家", llm_model=SimpleNamespace(model_name="integration-model"))
        data = {
            "title": self.current_item.title,
            "module_paths": [module.path],
            "description": self.current_item.description,
            "acceptance_criteria": "",
            "supplementary_description": "",
            "source_summary": "按新需求整合",
            "relationship_mode": "new",
            "change_type": "initial",
            "selected_revision_id": "",
            "conflicts": [],
            "open_questions": [],
        }

        RequirementReviewService._save_results(
            run,
            self.current_item,
            self.user,
            run.source_content_hash,
            "原始需求上下文",
            [],
            [],
            data,
            role,
        )

        run.refresh_from_db()
        draft = self.current_item.integration_draft
        self.assertEqual(run.status, "completed")
        self.assertEqual(draft.relationship_mode, "new")
        self.assertIsNone(draft.selected_family_id)

    def test_full_paths_match_same_named_modules_in_different_branches(self):
        root = ProjectModule.objects.create(project=self.project, code="root", name="项目根节点")
        web = ProjectModule.objects.create(project=self.project, parent=root, code="web", name="web")
        app = ProjectModule.objects.create(project=self.project, parent=root, code="app", name="app")
        web_login = ProjectModule.objects.create(project=self.project, parent=web, code="web-login", name="登录")
        app_login = ProjectModule.objects.create(project=self.project, parent=app, code="app-login", name="登录")

        suggested, matched, unresolved = RequirementReviewService._resolve_modules(
            self.current_item,
            ["项目根节点/web/登录", " 项目根节点 / app / 登录 ", "项目根节点/web/登录", "不存在 / 模块"],
        )

        self.assertEqual({module.id for module in matched}, {web_login.id, app_login.id})
        self.assertEqual(len(suggested), 3)
        self.assertEqual(unresolved, ["不存在 / 模块"])

    def test_manual_module_selection_resolves_unmatched_paths(self):
        module = ProjectModule.objects.create(project=self.project, code="auth", name="认证")
        draft = RequirementIntegrationDraft.objects.create(
            requirement_item=self.current_item,
            status="completed",
            suggested_module_paths=["项目根节点 / 认证"],
            unresolved_module_paths=["项目根节点 / 认证"],
            module_resolution_status="needs_review",
            created_by=self.user,
        )

        response = self.client.patch(
            reverse("requirement-item-integration", args=[self.current_item.id]),
            {"formal_module_ids": [module.id]},
            format="json",
        )

        self.assertEqual(response.status_code, 200)
        draft.refresh_from_db()
        self.assertEqual(list(draft.formal_modules.values_list("id", flat=True)), [module.id])
        self.assertEqual(draft.unresolved_module_paths, [])
        self.assertEqual(draft.module_resolution_status, "resolved")

    def test_formal_confirmation_persists_parallel_module_ids(self):
        root = ProjectModule.objects.create(project=self.project, code="root", name="项目根节点")
        web = ProjectModule.objects.create(project=self.project, parent=root, code="web", name="web")
        app = ProjectModule.objects.create(project=self.project, parent=root, code="app", name="app")
        source_hash = RequirementReviewService.source_hash(self.current_item)
        RequirementIntegrationRun.objects.create(
            requirement_item=self.current_item,
            status="completed",
            source_content_hash=source_hash,
            created_by=self.user,
        )
        draft = RequirementIntegrationDraft.objects.create(
            requirement_item=self.current_item,
            status="completed",
            review_status="approved",
            relationship_mode="new",
            relationship_confirmed=True,
            change_type="initial",
            module_resolution_status="resolved",
            source_content_hash=source_hash,
            title=self.current_item.title,
            description=self.current_item.description,
            created_by=self.user,
        )
        draft.formal_modules.set([web, app])

        revision = RequirementReviewService.confirm(self.current_item, self.user)

        self.assertEqual(set(revision.modules.values_list("id", flat=True)), {web.id, app.id})
        self.assertEqual(set(revision.family.modules.values_list("id", flat=True)), {web.id, app.id})
        self.assertEqual(set(self.current_item.formal_modules.values_list("id", flat=True)), {web.id, app.id})
        filtered = self.client.get(reverse("requirement-item-list"), {"formal_module": root.id})
        self.assertEqual(filtered.status_code, 200)
        self.assertIn(self.current_item.id, [item["id"] for item in filtered.data["results"]])

    def test_merge_preserves_all_source_module_labels(self):
        self.current_item.source_module_labels = ["Web 登录"]
        self.current_item.save(update_fields=["source_module_labels"])
        self.sibling_item.source_module_labels = ["账号中心", "Web 登录"]
        self.sibling_item.save(update_fields=["source_module_labels"])

        response = self.client.post(
            reverse("requirement-item-merge"),
            {"ids": [self.current_item.id, self.sibling_item.id], "title": "合并需求"},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data["source_module_labels"], ["Web 登录", "账号中心"])

    def test_formal_revision_is_not_assigned_when_draft_version_is_created(self):
        module = ProjectModule.objects.create(project=self.project, code="auth", name="认证")
        family = RequirementFamily.objects.create(
            project=self.project,
            family_no="FR-001",
            title="正式需求",
            created_by=self.user,
        )
        revision = RequirementRevision.objects.create(
            family=family,
            source_item=self.current_item,
            revision_no=1,
            change_type="initial",
            title="正式需求",
            description="审核通过后的正式需求",
            source_content_hash="a" * 64,
            confirmed_by=self.user,
        )
        family.modules.set([module])
        revision.modules.set([module])
        response = self.client.post(
            reverse("requirement-version-list"),
            {
                "project": self.project.id,
                "name": "首个版本",
                "version_no": "v1.0",
                "requirement_revisions": [revision.id],
            },
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        version = RequirementVersion.objects.get(pk=response.data["id"])
        self.assertEqual(version.status, "draft")
        self.assertFalse(version.requirement_revisions.exists())
        self.assertFalse(version.requirement_items.exists())
