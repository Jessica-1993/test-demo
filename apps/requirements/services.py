import base64
import os
import re
import tempfile
import json
import logging
import io
import zipfile
import ssl
from html import escape
from html.parser import HTMLParser
from pathlib import Path
from uuid import uuid4
from urllib.error import HTTPError, URLError
from urllib.parse import quote
from urllib.request import Request, urlopen

from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from django.utils import timezone

from apps.configuration.models import PromptConfig
from apps.configuration.services import normalize_gemini_base_url, normalize_gemini_model_name
from apps.core.errors import ClassifiedError, classify_message, error_info_from_exception, provider_error_info


logger = logging.getLogger(__name__)


class _SafeTableHTMLParser(HTMLParser):
    ALLOWED_TAGS = {"table", "thead", "tbody", "tr", "th", "td"}
    ALLOWED_ATTRIBUTES = {"rowspan", "colspan"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []

    def handle_starttag(self, tag, attrs):
        if tag not in self.ALLOWED_TAGS:
            return
        safe_attrs = "".join(
            f' {name}="{escape(value or "", quote=True)}"'
            for name, value in attrs if name in self.ALLOWED_ATTRIBUTES
        )
        self.parts.append(f"<{tag}{safe_attrs}>")

    def handle_endtag(self, tag):
        if tag in self.ALLOWED_TAGS:
            self.parts.append(f"</{tag}>")

    def handle_data(self, data):
        self.parts.append(escape(data))

    def html(self):
        return "".join(self.parts)


class QiniuStorageService:
    def __init__(self):
        self.access_key = getattr(settings, "QINIU_ACCESS_KEY", "")
        self.secret_key = getattr(settings, "QINIU_SECRET_KEY", "")
        self.bucket = getattr(settings, "QINIU_BUCKET", "")
        self.domain = self._normalize_domain(getattr(settings, "QINIU_DOMAIN", ""))
        self.prefix = getattr(settings, "QINIU_DOCS_PREFIX", "docs").strip("/")

    def upload(self, upload_file):
        if not all([self.access_key, self.secret_key, self.bucket]):
            raise ImproperlyConfigured("七牛云配置不完整，请配置 QINIU_ACCESS_KEY、QINIU_SECRET_KEY、QINIU_BUCKET")
        try:
            from qiniu import Auth, put_data
        except ImportError as exc:
            raise ImproperlyConfigured("缺少 qiniu 依赖，请先安装 requirements.txt") from exc

        suffix = Path(upload_file.name).suffix.lower()
        key = f"{self.prefix}/{uuid4().hex}{suffix}"
        auth = Auth(self.access_key, self.secret_key)
        token = auth.upload_token(self.bucket, key)
        ret, info = put_data(token, key, upload_file.read())
        if info.status_code != 200:
            raise RuntimeError(f"七牛上传失败: {info.text_body}")
        return {
            "key": ret.get("key", key),
            "url": f"{self.domain}/{key}" if self.domain else "",
        }

    def upload_bytes(self, content, key):
        if not all([self.access_key, self.secret_key, self.bucket]):
            raise ImproperlyConfigured("七牛云配置不完整")
        from qiniu import Auth, put_data
        auth = Auth(self.access_key, self.secret_key)
        ret, info = put_data(auth.upload_token(self.bucket, key), key, content)
        if info.status_code != 200:
            raise RuntimeError(f"七牛上传失败: {info.text_body}")
        return {"key": ret.get("key", key), "url": self.public_url(key)}

    def delete(self, key):
        key = (key or "").strip()
        if not key:
            return
        if not all([self.access_key, self.secret_key, self.bucket]):
            raise ImproperlyConfigured("七牛云配置不完整")
        try:
            from qiniu import Auth, BucketManager
        except ImportError as exc:
            raise ImproperlyConfigured("缺少 qiniu 依赖，请先安装 requirements.txt") from exc

        manager = BucketManager(Auth(self.access_key, self.secret_key))
        _ret, info = manager.delete(self.bucket, key)
        if info.status_code not in {200, 612}:
            raise RuntimeError(f"{key}: {info.text_body}")

    def delete_many(self, keys):
        unique_keys = []
        seen = set()
        for key in keys:
            normalized = (key or "").strip()
            if normalized and normalized not in seen:
                unique_keys.append(normalized)
                seen.add(normalized)

        failed = []
        for key in unique_keys:
            try:
                self.delete(key)
            except (ImproperlyConfigured, RuntimeError) as exc:
                failed.append(str(exc))
        if failed:
            raise RuntimeError("；".join(failed))

    def public_url(self, key):
        return f"{self.domain}/{quote(key, safe='/')}" if self.domain else ""

    def access_url(self, key, fallback_url=""):
        url = self.public_url(key) or self._normalize_domain(fallback_url)
        if not url or not self.access_key or not self.secret_key:
            return url
        try:
            from qiniu import Auth
        except ImportError:
            return url
        return Auth(self.access_key, self.secret_key).private_download_url(url, expires=3600)

    @staticmethod
    def _normalize_domain(domain):
        domain = (domain or "").strip().rstrip("/")
        if domain and not domain.startswith(("http://", "https://")):
            domain = f"https://{domain}"
        return domain

    def list_documents(self):
        if not all([self.access_key, self.secret_key, self.bucket]):
            raise ImproperlyConfigured("七牛云配置不完整")
        from qiniu import Auth, BucketManager
        manager = BucketManager(Auth(self.access_key, self.secret_key))
        marker = None
        objects = []
        while True:
            ret, _eof, info = manager.list(self.bucket, prefix=f"{self.prefix}/", marker=marker, limit=1000)
            if info.status_code != 200:
                raise RuntimeError(f"七牛文件列表读取失败: {info.text_body}")
            objects.extend(ret.get("items", []))
            marker = ret.get("marker")
            if not marker:
                break
        return objects

    def download(self, key, url=""):
        if not self.domain and not url:
            raise ImproperlyConfigured("重新解析需要配置 QINIU_DOMAIN 或保存可访问的七牛地址")
        download_url = f"{self.domain}/{quote(key, safe='/')}" if self.domain else url
        download_url = self._normalize_domain(download_url)
        if self.access_key and self.secret_key:
            try:
                from qiniu import Auth
            except ImportError as exc:
                raise ImproperlyConfigured("缺少 qiniu 依赖，请先安装 requirements.txt") from exc
            download_url = Auth(self.access_key, self.secret_key).private_download_url(download_url, expires=3600)
        try:
            with urlopen(Request(download_url), timeout=30) as remote_file:
                return remote_file.read()
        except URLError as exc:
            if isinstance(exc.reason, ssl.SSLCertVerificationError):
                raise RuntimeError(
                    "七牛文档下载失败：QINIU_DOMAIN 的 HTTPS 证书与域名不匹配，"
                    "请配置证书有效的 HTTPS 加速域名；如确认文档允许明文传输，可显式配置 http:// 域名"
                ) from exc
            raise RuntimeError(f"七牛文档下载失败: {exc}") from exc
        except (HTTPError, TimeoutError) as exc:
            raise RuntimeError(f"七牛文档下载失败: {exc}") from exc


class DocumentExtractionError(RuntimeError):
    pass


class DocumentExtractionService:
    PICTURE_NOTE = "图片内容暂未识别"

    @classmethod
    def extract_upload(cls, upload_file, document_type):
        upload_file.seek(0)
        suffix = Path(upload_file.name).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            for chunk in upload_file.chunks():
                temp.write(chunk)
            temp_path = temp.name
        try:
            return cls.extract_path(temp_path, document_type, upload_file.name)
        finally:
            os.unlink(temp_path)

    @classmethod
    def extract_bytes(cls, content, document_type, filename):
        suffix = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            temp.write(content)
            temp_path = temp.name
        try:
            return cls.extract_path(temp_path, document_type, filename)
        finally:
            os.unlink(temp_path)

    @classmethod
    def extract_path(cls, path, document_type, filename):
        if document_type in {"pdf", "docx"}:
            try:
                return cls._extract_with_docling(path, filename, document_type)
            except Exception as exc:
                logger.warning("Docling extraction failed for %s: %s", filename, exc)
                try:
                    result = cls._extract_legacy(path, filename, document_type)
                except Exception as fallback_exc:
                    raise DocumentExtractionError(str(fallback_exc)) from fallback_exc
                result["fallback_used"] = True
                return result
        try:
            text = RequirementParser._extract_text_file(path)
        except Exception as exc:
            raise DocumentExtractionError(str(exc)) from exc
        blocks = cls._text_blocks(text, filename, document_type)
        return cls._result(text, blocks, "text", False)

    @classmethod
    def _extract_with_docling(cls, path, filename, document_type):
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        format_options = None
        if document_type == "pdf":
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = False
            pipeline_options.do_table_structure = True
            pipeline_options.generate_picture_images = True
            format_options = {InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
        converter = DocumentConverter(format_options=format_options)
        doc = converter.convert(path).document
        blocks = []
        for item, tree_level in doc.iterate_items(traverse_pictures=False):
            label = cls._label_value(getattr(item, "label", ""))
            if label in {"page_header", "page_footer"}:
                text = (getattr(item, "text", "") or "").strip()
                cls._append_block(
                    blocks, filename, "paragraph", text, cls._page_number(item),
                    filtered_reason="页眉" if label == "page_header" else "页脚",
                )
                continue
            page = cls._page_number(item)
            if label == "table":
                markdown = item.export_to_markdown(doc=doc).strip()
                html = cls._sanitize_table_html(item.export_to_html(doc=doc).strip())
                cls._append_block(blocks, filename, "table", markdown, page, markdown=markdown, html=html, rows=cls._html_table_rows(html))
            elif label == "picture":
                caption = item.caption_text(doc).strip() if hasattr(item, "caption_text") else ""
                image_data = None
                image_width = image_height = None
                try:
                    image = item.get_image(doc)
                    if image:
                        image_width, image_height = image.size
                        output = io.BytesIO()
                        image.save(output, format="PNG")
                        image_data = output.getvalue()
                except Exception:
                    logger.debug("Docling picture export failed for %s", filename, exc_info=True)
                cls._append_block(
                    blocks, filename, "picture", caption or cls.PICTURE_NOTE, page,
                    image_data=image_data, image_width=image_width, image_height=image_height,
                )
            elif label in {"title", "section_header"}:
                text = getattr(item, "text", "").strip()
                if text:
                    level = getattr(item, "level", None) or max(1, min(tree_level + 1, 6))
                    cls._append_block(blocks, filename, "heading", text, page, level=level)
            elif hasattr(item, "text"):
                text = item.text.strip()
                if text:
                    cls._append_block(blocks, filename, "paragraph", text, page)
        blocks = cls._remove_duplicate_table_cells(blocks)
        plain_text = cls._blocks_to_text(blocks)
        if not plain_text:
            raise DocumentExtractionError("Docling 未提取到可用文本")
        return cls._result(plain_text, blocks, "docling", False)

    @classmethod
    def _extract_legacy(cls, path, filename, document_type):
        if document_type == "pdf":
            text, blocks = cls._legacy_pdf(path, filename)
            return cls._result(text, blocks, "legacy_pdf", True)
        text, blocks = cls._legacy_docx(path, filename)
        return cls._result(text, blocks, "legacy_docx", True)

    @classmethod
    def _legacy_pdf(cls, path, filename):
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise ImproperlyConfigured("解析 PDF 需要安装 PyPDF2") from exc
        blocks = []
        reader = PdfReader(path)
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            for paragraph in cls._split_paragraphs(page_text):
                block_type, level, text = cls._classify_text(paragraph)
                cls._append_block(blocks, filename, block_type, text, page_number, level=level)
            for _ in range(cls._count_pdf_images(page)):
                cls._append_block(blocks, filename, "picture", cls.PICTURE_NOTE, page_number)
        return cls._blocks_to_text(blocks), blocks

    @classmethod
    def _legacy_docx(cls, path, filename):
        try:
            import docx
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise ImproperlyConfigured("解析 Word 需要安装 python-docx") from exc
        document = docx.Document(path)
        blocks = []
        image_payloads = []
        with zipfile.ZipFile(path) as archive:
            for name in sorted(name for name in archive.namelist() if name.startswith("word/media/")):
                image_payloads.append(archive.read(name))
        image_index = 0
        for element in document.iter_inner_content():
            if isinstance(element, Paragraph):
                text = element.text.strip()
                style_name = (element.style.name or "") if element.style else ""
                if text:
                    heading_match = re.search(r"(\d+)$", style_name)
                    if style_name.lower().startswith("heading") or style_name.startswith("标题"):
                        level = int(heading_match.group(1)) if heading_match else 1
                        cls._append_block(blocks, filename, "heading", text, None, level=min(level, 6))
                    else:
                        cls._append_block(blocks, filename, "paragraph", text, None)
                drawing_count = len(element._element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))
                for _ in range(drawing_count):
                    image_data = image_payloads[image_index] if image_index < len(image_payloads) else None
                    image_index += 1
                    cls._append_block(blocks, filename, "picture", cls.PICTURE_NOTE, None, image_data=image_data)
            elif isinstance(element, Table):
                rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in element.rows]
                markdown = cls._table_to_markdown(rows)
                html = cls._table_to_html(rows)
                cls._append_block(blocks, filename, "table", markdown, None, markdown=markdown, html=html, rows=rows)
        return cls._blocks_to_text(blocks), blocks

    @classmethod
    def _text_blocks(cls, text, filename, document_type):
        blocks = []
        lines = text.splitlines()
        index = 0
        while index < len(lines):
            line = lines[index].strip()
            if not line:
                index += 1
                continue
            if document_type == "md" and line.startswith("|"):
                table_lines = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                markdown = "\n".join(table_lines)
                cls._append_block(blocks, filename, "table", markdown, None, markdown=markdown)
                continue
            block_type, level, normalized = cls._classify_text(line)
            cls._append_block(blocks, filename, block_type, normalized, None, level=level)
            index += 1
        return blocks

    @staticmethod
    def _label_value(label):
        value = getattr(label, "value", label)
        return str(value).lower()

    @staticmethod
    def _page_number(item):
        provenance = getattr(item, "prov", None) or []
        return getattr(provenance[0], "page_no", None) if provenance else None

    @classmethod
    def _append_block(cls, blocks, filename, block_type, text, page, **extra):
        block_index = len(blocks) + 1
        page_part = f":P{page}" if page else ""
        block = {
            "type": block_type,
            "text": text,
            "page": page,
            "block_index": block_index,
            "source_locator": f"{filename}{page_part}#B{block_index}",
        }
        block.update({key: value for key, value in extra.items() if value is not None})
        blocks.append(block)

    @staticmethod
    def _classify_text(text):
        markdown_heading = re.match(r"^(#{1,6})\s+(.+)$", text)
        if markdown_heading:
            return "heading", len(markdown_heading.group(1)), markdown_heading.group(2).strip()
        numeric_heading = re.match(r"^(\d+(?:[.．]\d+){0,5})[、.．\s]+(.+)$", text)
        if numeric_heading:
            return "heading", min(len(re.split(r"[.．]", numeric_heading.group(1))), 6), numeric_heading.group(2).strip()
        heading = RequirementParser.HEADING_PATTERN.match(text)
        if heading:
            return "heading", 1, heading.group(2).strip()
        return "paragraph", None, text

    @staticmethod
    def _split_paragraphs(text):
        chunks = re.split(r"\n\s*\n", text)
        return [line.strip() for chunk in chunks for line in chunk.splitlines() if line.strip()]

    @staticmethod
    def _count_pdf_images(page):
        try:
            resources = page.get("/Resources") or {}
            xobjects = resources.get("/XObject") or {}
            return sum(1 for obj in xobjects.values() if obj.get_object().get("/Subtype") == "/Image")
        except (AttributeError, KeyError, TypeError):
            return 0

    @staticmethod
    def _table_to_markdown(rows):
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = normalized[0]
        body = normalized[1:]
        render = lambda row: "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"
        return "\n".join([render(header), render(["---"] * width), *(render(row) for row in body)])

    @staticmethod
    def _table_to_html(rows):
        if not rows:
            return ""
        rendered = []
        for row_index, row in enumerate(rows):
            tag = "th" if row_index == 0 else "td"
            rendered.append("<tr>" + "".join(f"<{tag}>{escape(cell)}</{tag}>" for cell in row) + "</tr>")
        return "<table>" + "".join(rendered) + "</table>"

    @staticmethod
    def _html_table_rows(html):
        rows = []
        for row in re.findall(r"<tr[^>]*>(.*?)</tr>", html, flags=re.I | re.S):
            cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row, flags=re.I | re.S)
            rows.append([re.sub(r"<[^>]+>", "", cell).strip() for cell in cells])
        return rows

    @staticmethod
    def _remove_duplicate_table_cells(blocks):
        cleaned = []
        table_cells = set()
        checking_table_children = False
        for block in blocks:
            if block.get("type") == "table":
                cleaned.append(block)
                table_cells = {
                    re.sub(r"\s+", " ", str(cell)).strip()
                    for row in block.get("rows", []) for cell in row if str(cell).strip()
                }
                checking_table_children = bool(table_cells)
                continue
            if checking_table_children and block.get("type") == "paragraph":
                text = re.sub(r"\s+", " ", block.get("text", "")).strip()
                if text in table_cells:
                    continue
            checking_table_children = False
            table_cells = set()
            cleaned.append(block)
        return cleaned

    @staticmethod
    def _sanitize_table_html(html):
        parser = _SafeTableHTMLParser()
        parser.feed(html)
        return parser.html()

    @staticmethod
    def _blocks_to_text(blocks):
        parts = []
        for block in blocks:
            if block["type"] == "picture":
                continue
            if block["type"] == "heading":
                parts.append(f"{'#' * block.get('level', 1)} {block['text']}")
            else:
                parts.append(block["text"])
        return "\n\n".join(part for part in parts if part).strip()

    @staticmethod
    def _result(text, blocks, engine, fallback_used):
        return {
            "plain_text": text.strip(),
            "blocks": blocks,
            "engine": engine,
            "fallback_used": fallback_used,
        }


class RequirementParser:
    HEADING_PATTERN = re.compile(r"^(#{1,6}\s+|[一二三四五六七八九十]+[、.．]|\d+[.．、])\s*(.+)$")

    @classmethod
    def extract_text(cls, upload_file, document_type):
        upload_file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(upload_file.name).suffix.lower()) as temp:
            for chunk in upload_file.chunks():
                temp.write(chunk)
            temp_path = temp.name
        try:
            if document_type == "pdf":
                return cls._extract_pdf(temp_path)
            if document_type == "docx":
                return cls._extract_docx(temp_path)
            return cls._extract_text_file(temp_path)
        finally:
            os.unlink(temp_path)

    @staticmethod
    def detect_type(filename):
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        if suffix in [".doc", ".docx"]:
            return "docx"
        if suffix == ".md":
            return "md"
        if suffix == ".txt":
            return "txt"
        return "other"

    @staticmethod
    def _extract_text_file(path):
        for encoding in ["utf-8", "gbk"]:
            try:
                return Path(path).read_text(encoding=encoding).strip()
            except UnicodeDecodeError:
                continue
        return Path(path).read_text(errors="ignore").strip()

    @staticmethod
    def _extract_pdf(path):
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise ImproperlyConfigured("解析 PDF 需要安装 PyPDF2") from exc
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    @staticmethod
    def _extract_docx(path):
        try:
            import docx
        except ImportError as exc:
            raise ImproperlyConfigured("解析 Word 需要安装 python-docx") from exc
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs if p.text.strip()).strip()


class StructuredRequirementParser:
    EXCLUDED_HEADINGS = {
        "封面", "目录", "修订记录", "变更记录", "版本记录", "术语", "术语解释",
        "项目背景", "背景", "项目目标", "目标", "参考资料", "参考文档",
    }
    PLACEHOLDER_PATTERNS = ("此处插入图片", "图片占位", "插入图片", "图片待补充")

    @classmethod
    def organize(cls, blocks, default_title):
        document_title = default_title
        module = "未分类"
        current = None
        requirements = []
        orphan_blocks = []
        filtered = []
        excluded_level = None

        def finish_current():
            nonlocal current
            if not current:
                return
            if current["blocks"]:
                requirements.append(current)
            else:
                filtered.append({"source_locator": current["source_locator"], "reason": "只有标题但没有有效正文"})
            current = None

        for raw in blocks:
            block = dict(raw)
            block_type = block.get("type")
            level = int(block.get("level") or 0) if block_type == "heading" else 0
            text = (block.get("text") or "").strip()

            if block.get("filtered_reason"):
                filtered.append({
                    "source_locator": block.get("source_locator", ""),
                    "reason": block["filtered_reason"],
                })
                continue

            if excluded_level is not None:
                if block_type == "heading" and level <= excluded_level:
                    excluded_level = None
                else:
                    filtered.append({"source_locator": block.get("source_locator", ""), "reason": "默认排除章节"})
                    continue

            if block_type == "heading" and cls._normalized_heading(text) in cls.EXCLUDED_HEADINGS:
                finish_current()
                excluded_level = level or 1
                filtered.append({"source_locator": block.get("source_locator", ""), "reason": f"默认排除章节：{text}"})
                continue

            if block_type == "paragraph" and any(pattern in text for pattern in cls.PLACEHOLDER_PATTERNS):
                filtered.append({"source_locator": block.get("source_locator", ""), "reason": "图片占位说明"})
                continue

            if block_type == "picture" and not block.get("image_data"):
                filtered.append({"source_locator": block.get("source_locator", ""), "reason": "未提取到真实图片"})
                continue

            if block_type == "heading":
                text = cls._clean_heading(text)
                if level == 1:
                    finish_current()
                    document_title = text or document_title
                elif level == 2:
                    finish_current()
                    module = text or "未分类"
                elif level == 3:
                    finish_current()
                    current = {
                        "title": text or "未命名需求", "module": module, "description": [],
                        "supplementary": [], "blocks": [], "source_locator": block.get("source_locator", ""),
                    }
                elif level >= 4 and current:
                    current["supplementary"].append(text)
                    current["blocks"].append(cls._content_block(block, "text", text, level))
                continue

            content = cls._content_block(block)
            if current:
                current["blocks"].append(content)
                if block_type == "paragraph":
                    current["description"].append(text)
            else:
                orphan_blocks.append(content)

        finish_current()
        return {
            "document_title": document_title,
            "requirements": requirements,
            "orphan_blocks": orphan_blocks,
            "filtered": filtered,
        }

    @staticmethod
    def _normalized_heading(text):
        text = re.sub(r"^第?[一二三四五六七八九十百0-9]+[章节部分、.．\s]+", "", text.strip())
        return re.sub(r"[：:\s]", "", text)

    @staticmethod
    def _clean_heading(text):
        prefix = r"(?:REQ[-_]\d+|第?[一二三四五六七八九十百0-9]+[章节部分]|\d+(?:[.．]\d+)*)"
        cleaned = re.sub(rf"^\s*{prefix}[、.．：:\s]+", "", text.strip())
        return cleaned.strip() or text.strip()

    @staticmethod
    def _content_block(block, forced_type=None, forced_text=None, heading_level=None):
        type_map = {"paragraph": "text", "table": "table", "picture": "image"}
        return {
            "block_type": forced_type or type_map.get(block.get("type"), "text"),
            "text": forced_text if forced_text is not None else block.get("text", ""),
            "heading_level": heading_level,
            "page": block.get("page"),
            "source_locator": block.get("source_locator", ""),
            "table_data": {
                "rows": block.get("rows", []),
                "html": block.get("html", ""),
                "markdown": block.get("markdown", ""),
            } if block.get("type") == "table" else {},
            "image_data": block.get("image_data"),
            "image_width": block.get("image_width"),
            "image_height": block.get("image_height"),
        }

class TestCaseGenerationError(ClassifiedError):
    default_code = "MODEL_RESPONSE_INVALID"

    def __init__(self, message="", *, code=None, details=None, cause_detail="", trace_id=None):
        inferred = code or classify_message(message)
        if inferred == "VALIDATION_ERROR":
            inferred = self.default_code
        super().__init__(
            message, code=inferred, details=details,
            cause_detail=cause_detail, trace_id=trace_id,
        )


class VisionAnalysisError(TestCaseGenerationError):
    pass


class LLMChatService:
    def __init__(self, config):
        self.config = config

    def chat(self, system_prompt, user_prompt):
        if self.config.protocol == "gemini":
            return self._chat_gemini(system_prompt, user_prompt)
        return self._chat_openai_compatible(system_prompt, user_prompt)

    def _chat_openai_compatible(self, system_prompt, user_prompt):
        base_url = self.config.base_url.rstrip("/")
        if base_url.endswith("/v1/chat/completions"):
            url = base_url
        elif base_url.endswith("/v1"):
            url = f"{base_url}/chat/completions"
        else:
            url = f"{base_url}/v1/chat/completions"
        payload = {
            "model": self.config.model_name,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "max_tokens": self.config.max_tokens,
            "temperature": self.config.temperature,
            "top_p": self.config.top_p,
            "stream": False,
        }
        headers = {
            "Authorization": f"Bearer {self.config.api_key}",
            "Content-Type": "application/json",
        }
        data = self._post_json(url, payload, headers, provider=self.config.provider, stage="模型生成")
        try:
            return data["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError) as exc:
            raise TestCaseGenerationError("模型响应缺少 choices[0].message.content") from exc

    def _chat_gemini(self, system_prompt, user_prompt):
        base_url = normalize_gemini_base_url(self.config.base_url)
        model = quote(normalize_gemini_model_name(self.config.model_name), safe="/")
        url = f"{base_url}/{model}:generateContent?key={quote(self.config.api_key, safe='')}"
        payload = {
            "systemInstruction": {"parts": [{"text": system_prompt}]},
            "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
            "generationConfig": {
                "maxOutputTokens": self.config.max_tokens,
                "temperature": self.config.temperature,
                "topP": self.config.top_p,
            },
        }
        data = self._post_json(
            url, payload, {"Content-Type": "application/json"},
            provider=self.config.provider, stage="模型生成",
        )
        try:
            return "".join(part.get("text", "") for part in data["candidates"][0]["content"]["parts"]).strip()
        except (KeyError, IndexError, TypeError) as exc:
            raise TestCaseGenerationError("Gemini 响应缺少 candidates[0].content.parts") from exc

    @staticmethod
    def _post_json(url, payload, headers, *, provider="", stage="模型调用"):
        request = Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers, method="POST")
        try:
            with urlopen(request, timeout=60) as res:
                return json.loads(res.read().decode("utf-8") or "{}")
        except HTTPError as exc:
            body = exc.read().decode("utf-8", errors="replace")
            info, diagnostic = provider_error_info(exc.code, body, provider=provider, stage=stage)
            raise TestCaseGenerationError(
                info["message"], code=info["code"], details=info["details"],
                cause_detail=diagnostic, trace_id=info["trace_id"],
            ) from exc
        except URLError as exc:
            raise TestCaseGenerationError(
                "模型服务暂时不可用", code="MODEL_PROVIDER_UNAVAILABLE",
                details={"provider": provider, "stage": stage}, cause_detail=str(exc.reason),
            ) from exc
        except TimeoutError as exc:
            raise TestCaseGenerationError(
                "模型调用超时", code="MODEL_TIMEOUT",
                details={"provider": provider, "stage": stage},
            ) from exc
        except json.JSONDecodeError as exc:
            raise TestCaseGenerationError(
                "模型响应不是合法 JSON", code="MODEL_RESPONSE_INVALID",
                details={"provider": provider, "stage": stage},
            ) from exc


class LLMResponsesVisionService:
    def __init__(self, config):
        self.config = config

    def analyze_image(self, image_url, instruction):
        if self.config.protocol != "openai_responses":
            raise VisionAnalysisError("图片理解模型必须使用 OpenAI Responses 协议")
        base_url = self.config.base_url.rstrip("/")
        if base_url.endswith("/v1/responses"):
            url = base_url
        elif base_url.endswith("/v1"):
            url = f"{base_url}/responses"
        else:
            url = f"{base_url}/v1/responses"
        payload = {
            "model": self.config.model_name,
            "input": [
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": instruction},
                        {"type": "input_image", "image_url": image_url},
                    ],
                }
            ],
            "max_output_tokens": self.config.max_tokens,
            "temperature": self.config.temperature,
            "top_p": self.config.top_p,
        }
        headers = {
            "Authorization": f"Bearer {self.config.api_key}",
            "Content-Type": "application/json",
        }
        data = LLMChatService._post_json(
            url, payload, headers, provider=self.config.provider, stage="图片理解",
        )
        output_text = self._extract_output_text(data)
        if not output_text:
            raise VisionAnalysisError("图片理解模型未返回文本摘要")
        return output_text, data

    @staticmethod
    def _extract_output_text(data):
        if isinstance(data.get("output_text"), str):
            return data["output_text"].strip()
        parts = []
        for item in data.get("output", []):
            for content in item.get("content", []):
                text = content.get("text")
                if text:
                    parts.append(text)
        return "\n".join(parts).strip()


class LLMGeminiVisionService:
    def __init__(self, config):
        self.config = config

    def analyze_image(self, image_url, instruction):
        if self.config.protocol != "gemini":
            raise VisionAnalysisError("Gemini 图片理解模型必须使用 Gemini 协议")
        image_bytes, mime_type = self._download_image(image_url)
        base_url = normalize_gemini_base_url(self.config.base_url)
        model = quote(normalize_gemini_model_name(self.config.model_name), safe="/")
        url = f"{base_url}/{model}:generateContent?key={quote(self.config.api_key, safe='')}"
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        {"text": instruction},
                        {
                            "inline_data": {
                                "mime_type": mime_type,
                                "data": base64.b64encode(image_bytes).decode("ascii"),
                            }
                        },
                    ],
                }
            ],
            "generationConfig": {
                "maxOutputTokens": self.config.max_tokens,
                "temperature": self.config.temperature,
                "topP": self.config.top_p,
            },
        }
        data = LLMChatService._post_json(
            url, payload, {"Content-Type": "application/json"},
            provider=self.config.provider, stage="图片理解",
        )
        output_text = self._extract_output_text(data)
        if not output_text:
            raise VisionAnalysisError("Gemini 图片理解模型未返回文本摘要")
        return output_text, data

    @staticmethod
    def _download_image(image_url):
        try:
            with urlopen(Request(image_url), timeout=30) as response:
                content = response.read()
                content_type = response.headers.get("Content-Type", "").split(";", 1)[0].strip()
        except (HTTPError, URLError, TimeoutError) as exc:
            raise VisionAnalysisError(f"图片下载失败: {exc}") from exc
        if not content:
            raise VisionAnalysisError("图片下载结果为空")
        return content, content_type or "image/png"

    @staticmethod
    def _extract_output_text(data):
        try:
            return "".join(part.get("text", "") for part in data["candidates"][0]["content"]["parts"]).strip()
        except (KeyError, IndexError, TypeError) as exc:
            raise VisionAnalysisError("Gemini 响应缺少 candidates[0].content.parts") from exc


class RequirementImageAnalysisService:
    SUMMARY_FIELDS = [
        "image_type",
        "ui_elements",
        "business_rules",
        "flows",
        "states",
        "test_points",
        "uncertainties",
    ]

    @classmethod
    def ensure_for_requirement(cls, requirement_item):
        image_blocks = [
            block for block in requirement_item.content_blocks.all()
            if block.block_type == "image" and (block.image_key or block.image_url)
        ]
        if not image_blocks:
            return []
        role = cls.get_vision_role()
        results = []
        for block in image_blocks:
            analysis = cls.ensure_for_block(block, role)
            if analysis.status != "completed":
                raise VisionAnalysisError(f"图片内容识别失败: {analysis.error_message or block.source_locator or block.id}")
            results.append(analysis)
        return results

    @staticmethod
    def get_vision_role():
        role = PromptConfig.resolve_active("vision_analyzer", error_class=VisionAnalysisError)
        config = role.llm_model
        if config.protocol not in {"openai_responses", "gemini"}:
            raise VisionAnalysisError("图片理解模型仅支持 OpenAI Responses 或 Gemini 协议")
        return role

    @classmethod
    def get_vision_config(cls):
        return cls.get_vision_role().llm_model

    @classmethod
    def ensure_for_block(cls, block, role):
        from .models import RequirementImageAnalysis

        config = role.llm_model
        analysis, _created = RequirementImageAnalysis.objects.get_or_create(content_block=block)
        if analysis.status == "completed" and analysis.summary:
            return analysis
        image_url = QiniuStorageService().access_url(block.image_key, block.image_url)
        if not image_url:
            info = error_info_from_exception(
                VisionAnalysisError("图片没有可访问地址", code="STORAGE_UNAVAILABLE"),
                details={"stage": "图片理解", "resource": str(block.id)},
            )
            analysis.status = "failed"
            analysis.model_name = config.model_name
            analysis.error_message = info["message"]
            analysis.error_info = info
            analysis.analyzed_at = timezone.now()
            analysis.save(update_fields=["status", "model_name", "error_message", "error_info", "analyzed_at", "updated_at"])
            return analysis
        try:
            vision_service = LLMGeminiVisionService(config) if config.protocol == "gemini" else LLMResponsesVisionService(config)
            output_text, raw_response = vision_service.analyze_image(
                image_url,
                f"{role.prompt_content.strip()}\n\n{cls._build_instruction(block)}".strip(),
            )
            summary = cls._parse_summary(output_text)
            analysis.status = "completed"
            analysis.model_name = config.model_name
            analysis.summary = summary
            analysis.raw_response = cls._safe_raw_response(raw_response, output_text)
            analysis.error_message = ""
            analysis.error_info = {}
        except Exception as exc:
            info = error_info_from_exception(
                exc, details={"stage": "图片理解", "resource": str(block.id)},
            )
            analysis.status = "failed"
            analysis.model_name = config.model_name
            analysis.error_message = info["message"]
            analysis.error_info = info
        analysis.analyzed_at = timezone.now()
        analysis.save(update_fields=[
            "status", "model_name", "summary", "raw_response", "error_message", "error_info", "analyzed_at", "updated_at",
        ])
        return analysis

    @classmethod
    def _build_instruction(cls, block):
        locator = block.source_locator or f"内容块 {block.id}"
        return (
            "你是测试需求图片理解助手。请识别需求文档中的图片，输出严格 JSON 对象，不要 Markdown。\n"
            "JSON 字段必须包含: image_type, ui_elements, business_rules, flows, states, test_points, uncertainties。\n"
            "字段值可以是字符串数组或对象数组；无法确认的信息写入 uncertainties，不要臆造。\n"
            f"图片来源: {locator}\n"
            f"图片说明: {block.text or '无'}"
        )

    @classmethod
    def _parse_summary(cls, output_text):
        cleaned = TestCaseGenerationService._strip_code_fence(output_text)
        candidates = [cleaned]
        object_match = re.search(r"\{[\s\S]*\}", cleaned)
        if object_match:
            candidates.append(object_match.group(0))
        for candidate in candidates:
            try:
                data = json.loads(candidate)
            except json.JSONDecodeError:
                continue
            if isinstance(data, dict):
                return {field: data.get(field, [] if field != "image_type" else "") for field in cls.SUMMARY_FIELDS}
        raise VisionAnalysisError(
            "图片理解模型返回内容无法处理",
            code="MODEL_RESPONSE_INVALID",
            cause_detail=f"response_preview={TestCaseGenerationService._preview_content(output_text)}",
        )

    @staticmethod
    def _safe_raw_response(raw_response, output_text):
        if isinstance(raw_response, dict):
            return {
                "id": raw_response.get("id", ""),
                "model": raw_response.get("model", ""),
                "output_text": output_text,
            }
        return {"output_text": output_text}


class RequirementContextBuilder:
    @classmethod
    def build(cls, requirement_item):
        blocks = sorted(requirement_item.content_blocks.all(), key=lambda block: (block.order, block.id))
        parts = []
        for block in blocks:
            locator = block.source_locator or f"内容块 {block.id}"
            if block.block_type == "text" and block.text.strip():
                parts.append(f"[文本][{locator}]\n{block.text.strip()}")
            elif block.block_type == "table":
                table_text = cls._format_table(block.table_data)
                if table_text:
                    parts.append(f"[表格][{locator}]\n{table_text}")
            elif block.block_type == "image":
                parts.append(f"[图片][{locator}]\n{cls._format_image(block)}")
        block_context = "\n\n".join(parts).strip()
        return (
            f"需求编号: {requirement_item.requirement_no}\n"
            f"需求标题: {requirement_item.title}\n"
            f"功能模块: {requirement_item.module}\n"
            f"优先级: {requirement_item.get_priority_display()}\n"
            f"需求描述:\n{requirement_item.description}\n"
            f"验收标准:\n{requirement_item.acceptance_criteria or '无'}\n"
            f"补充描述:\n{requirement_item.supplementary_description or '无'}\n"
            f"原文结构化内容:\n{block_context or '无'}\n"
        )

    @staticmethod
    def _format_table(table_data):
        table_data = table_data or {}
        if table_data.get("markdown"):
            return table_data["markdown"]
        rows = table_data.get("rows") or []
        return DocumentExtractionService._table_to_markdown(rows) if rows else ""

    @staticmethod
    def _format_image(block):
        from .models import RequirementImageAnalysis

        analysis = RequirementImageAnalysis.objects.filter(content_block=block).first()
        if not analysis or analysis.status != "completed":
            return "图片内容尚未完成识别"
        return json.dumps(analysis.summary, ensure_ascii=False, indent=2)


class RequirementIntegrationService:
    EDITABLE_FIELDS = [
        "title",
        "module",
        "description",
        "acceptance_criteria",
        "supplementary_description",
        "source_summary",
    ]

    @classmethod
    def integrate(cls, requirement_item, user=None):
        from .models import RequirementIntegrationDraft

        role = cls.get_active_role()
        model = role.llm_model
        raw_context = RequirementContextBuilder.build(requirement_item)
        draft, created = RequirementIntegrationDraft.objects.get_or_create(
            requirement_item=requirement_item,
            defaults={"created_by": user or requirement_item.confirmed_by},
        )
        if user and not draft.created_by_id:
            draft.created_by = user
        draft.status = "pending"
        draft.raw_context = raw_context
        draft.model_name = model.model_name
        draft.prompt_name = role.name
        draft.error_message = ""
        draft.error_info = {}
        if user:
            draft.updated_by = user
        draft.save()
        try:
            content = LLMChatService(model).chat(role.prompt_content, cls._build_user_prompt(requirement_item, raw_context))
            data = cls.parse_integration(content)
        except Exception as exc:
            info = error_info_from_exception(
                exc, details={"stage": "需求整合", "resource": str(requirement_item.id)},
            )
            draft.status = "failed"
            draft.error_message = info["message"]
            draft.error_info = info
            draft.save(update_fields=["status", "error_message", "error_info", "updated_at"])
            raise
        for field in cls.EDITABLE_FIELDS:
            setattr(draft, field, str(data.get(field) or getattr(requirement_item, field, "") or ""))
        draft.status = "completed"
        draft.error_message = ""
        draft.error_info = {}
        draft.save()
        return draft

    @staticmethod
    def get_active_role():
        return PromptConfig.resolve_active("requirement_integrator", error_class=TestCaseGenerationError)

    @staticmethod
    def _build_user_prompt(requirement_item, raw_context):
        return (
            "请将下面单条详细需求整理为可直接生成测试用例的需求整合稿。\n"
            "输出必须是合法 JSON 对象，不要输出 Markdown 代码块、解释、标题或结尾说明。\n"
            "JSON 字段必须包含: title, module, description, acceptance_criteria, supplementary_description, source_summary。\n"
            "description 应整合文本、表格规则和图片识别结论；source_summary 用于说明依据和不确定点。\n\n"
            f"详细需求 ID: {requirement_item.id}\n"
            f"原始上下文:\n{raw_context}"
        )

    @classmethod
    def parse_integration(cls, content):
        cleaned = TestCaseGenerationService._strip_code_fence(content or "")
        candidates = [cleaned]
        object_match = re.search(r"\{[\s\S]*\}", cleaned)
        if object_match:
            candidates.append(object_match.group(0))
        for candidate in candidates:
            try:
                data = json.loads(candidate)
            except json.JSONDecodeError:
                continue
            if isinstance(data, dict):
                return {field: data.get(field, "") for field in cls.EDITABLE_FIELDS}
        preview = TestCaseGenerationService._preview_content(content)
        raise TestCaseGenerationError(
            "需求整合模型返回内容无法处理",
            code="MODEL_RESPONSE_INVALID",
            cause_detail=f"response_preview={preview}",
        )

    @staticmethod
    def build_generation_context(requirement_item, draft):
        module_paths = [module.path for module in draft.formal_modules.all()]
        return (
            f"需求编号: {requirement_item.requirement_no}\n"
            f"需求标题: {draft.title or requirement_item.title}\n"
            f"正式模块: {'；'.join(module_paths) or '未设置'}\n"
            f"原始模块标签: {'；'.join(requirement_item.source_module_labels or [requirement_item.module])}\n"
            f"优先级: {requirement_item.get_priority_display()}\n"
            f"需求描述:\n{draft.description or requirement_item.description}\n"
            f"验收标准:\n{draft.acceptance_criteria or '无'}\n"
            f"补充描述:\n{draft.supplementary_description or '无'}\n"
            f"来源摘要:\n{draft.source_summary or '无'}\n"
            f"原始上下文:\n{draft.raw_context or '无'}\n"
        )


class TestCaseGenerationService:
    REVIEW_PASS_KEYWORDS = ["通过", "合格", "合理", "完善", "无需修改", "pass", "approved"]
    REVIEW_FAIL_KEYWORDS = ["不通过", "不合格", "不合理", "不完善", "缺少", "补充", "修改", "重写", "fail", "rejected"]
    CASE_BATCH_SIZE = 6
    MAX_GENERATION_ROUNDS = 5

    @classmethod
    def get_active_role(cls, role_type):
        return PromptConfig.resolve_active(role_type, error_class=TestCaseGenerationError)

    @classmethod
    def generate_for_requirement(cls, requirement_item, review_feedback="", requirement_context=None):
        cases, content, writer_model, writer_role, _rounds = cls.generate_all_for_requirement(
            requirement_item,
            review_feedback=review_feedback,
            requirement_context=requirement_context,
        )
        return cases, content, writer_model, writer_role

    @classmethod
    def generate_all_for_requirement(cls, requirement_item, review_feedback="", requirement_context=None):
        writer_role = cls.get_active_role("testcase_writer")
        writer_model = writer_role.llm_model
        all_cases = []
        raw_contents = []
        seen_keys = set()
        for round_no in range(1, cls.MAX_GENERATION_ROUNDS + 1):
            user_prompt = cls._build_writer_user_prompt(
                requirement_item,
                review_feedback,
                requirement_context,
                existing_cases=all_cases,
                round_no=round_no,
            )
            content = LLMChatService(writer_model).chat(writer_role.prompt_content, user_prompt)
            raw_contents.append(content)
            cases = cls.parse_test_cases(content)
            if not cases:
                if all_cases:
                    break
                preview = cls._preview_content(content)
                raise TestCaseGenerationError(
                    "模型返回内容无法处理",
                    code="MODEL_RESPONSE_INVALID",
                    cause_detail=f"response_preview={preview}",
                )
            new_cases = cls._deduplicate_cases(cases, seen_keys)
            if not new_cases:
                break
            all_cases.extend(new_cases)
            if len(cases) < cls.CASE_BATCH_SIZE:
                break
        if not all_cases:
            raise TestCaseGenerationError("模型未返回可解析的测试用例")
        return all_cases, "\n\n".join(raw_contents), writer_model, writer_role, len(raw_contents)

    @classmethod
    def review_cases(cls, requirement_item, cases, requirement_context=None):
        reviewer_role = cls.get_active_role("testcase_reviewer")
        reviewer_model = reviewer_role.llm_model
        user_prompt = (
            "请审核下面测试用例是否覆盖该详细需求，返回明确结论。\n\n"
            f"详细需求:\n{cls._format_requirement(requirement_item, requirement_context)}\n\n"
            f"测试用例 JSON:\n{json.dumps(cases, ensure_ascii=False, indent=2)}"
        )
        feedback = LLMChatService(reviewer_model).chat(reviewer_role.prompt_content, user_prompt)
        return cls.review_passed(feedback), feedback, reviewer_model, reviewer_role

    @classmethod
    def review_passed(cls, feedback):
        text = (feedback or "").lower()
        if any(keyword in text for keyword in cls.REVIEW_FAIL_KEYWORDS):
            return False
        if any(keyword in text for keyword in cls.REVIEW_PASS_KEYWORDS):
            return True
        return True

    @classmethod
    def _build_writer_user_prompt(cls, requirement_item, review_feedback="", requirement_context=None, existing_cases=None, round_no=1):
        existing_cases = existing_cases or []
        existing_part = ""
        if existing_cases:
            existing_summary = [
                {
                    "case_no": case["case_no"],
                    "title": case["title"],
                    "test_type": case["test_type"],
                    "priority": case["priority"],
                }
                for case in existing_cases
            ]
            existing_part = (
                "\n\n已生成用例摘要:\n"
                f"{json.dumps(existing_summary, ensure_ascii=False)}\n"
                "请只补充尚未覆盖的新场景，不要重复已生成用例。若没有更多有价值场景，请返回 []。"
            )
        prompt = (
            "请针对单个详细需求生成测试用例。\n"
            "输出必须是合法 JSON 数组，不要输出 Markdown 代码块、表格、解释、标题或结尾说明。\n"
            "每个数组对象必须包含字段: case_no, title, preconditions, steps, expected_result, priority, test_type。\n"
            "priority 只允许 high/medium/low，test_type 只允许 functional/api/ui/integration/performance/security。\n\n"
            f"这是第 {round_no} 轮生成。本轮最多输出 {cls.CASE_BATCH_SIZE} 条测试用例；steps 和 expected_result 保持精简，避免输出过长导致 JSON 被截断。\n"
            "必须覆盖结构化内容中的文本、表格规则和图片识别测试点；图片存在 uncertainties 时，用例中只覆盖可确认信息。\n\n"
            "输出示例:\n"
            "[{\"case_no\":\"TC-001\",\"title\":\"验证用户登录成功\",\"preconditions\":\"存在有效用户账号\","
            "\"steps\":\"1. 打开登录页\\n2. 输入有效账号密码\\n3. 点击登录\","
            "\"expected_result\":\"登录成功并进入工作台\",\"priority\":\"high\",\"test_type\":\"functional\"}]\n\n"
            f"详细需求:\n{cls._format_requirement(requirement_item, requirement_context)}"
            f"{existing_part}"
        )
        if review_feedback:
            prompt += f"\n\n上一轮审核意见:\n{review_feedback}\n请根据审核意见重新生成完整用例。"
        return prompt

    @staticmethod
    def _preview_content(content, limit=500):
        text = re.sub(r"\s+", " ", content or "").strip()
        if not text:
            return "空响应"
        return text[:limit]

    @staticmethod
    def _format_requirement(requirement_item, requirement_context=None):
        return requirement_context or RequirementContextBuilder.build(requirement_item)

    @classmethod
    def parse_test_cases(cls, content):
        cleaned = cls._strip_code_fence(content or "")
        for candidate in cls._json_candidates(cleaned):
            try:
                data = json.loads(candidate)
            except json.JSONDecodeError:
                continue
            if isinstance(data, dict):
                wrapped = data.get("test_cases") or data.get("cases")
                if wrapped is None:
                    continue
                data = wrapped
            if isinstance(data, list):
                return [cls._normalize_case(item, index) for index, item in enumerate(data, start=1) if isinstance(item, dict)]
        partial_cases = cls._parse_partial_json_array(cleaned)
        if partial_cases:
            return partial_cases
        return cls._parse_markdown_table(cleaned)

    @staticmethod
    def _strip_code_fence(content):
        return re.sub(r"^```(?:json)?\s*|\s*```$", "", content.strip(), flags=re.IGNORECASE | re.MULTILINE).strip()

    @staticmethod
    def _json_candidates(content):
        candidates = [content]
        array_match = re.search(r"\[[\s\S]*\]", content)
        if array_match:
            candidates.append(array_match.group(0))
        object_match = re.search(r"\{[\s\S]*\}", content)
        if object_match:
            candidates.append(object_match.group(0))
        return candidates

    @classmethod
    def _parse_partial_json_array(cls, content):
        start = content.find("[")
        if start < 0:
            return []
        decoder = json.JSONDecoder()
        index = start + 1
        items = []
        while index < len(content):
            while index < len(content) and content[index] in " \r\n\t,":
                index += 1
            if index >= len(content) or content[index] == "]":
                break
            try:
                item, end = decoder.raw_decode(content, index)
            except json.JSONDecodeError:
                break
            if isinstance(item, dict):
                items.append(item)
            index = end
        return [cls._normalize_case(item, index) for index, item in enumerate(items, start=1)]

    @classmethod
    def _normalize_case(cls, item, index):
        priority = item.get("priority") or item.get("优先级") or "medium"
        test_type = item.get("test_type") or item.get("测试类型") or "functional"
        return {
            "case_no": str(item.get("case_no") or item.get("case_id") or item.get("用例ID") or f"TC-{index:03d}")[:80],
            "title": str(item.get("title") or item.get("测试目标") or item.get("用例标题") or f"测试用例 {index}")[:300],
            "preconditions": str(item.get("preconditions") or item.get("precondition") or item.get("前置条件") or ""),
            "steps": cls._stringify_steps(item.get("steps") or item.get("test_steps") or item.get("操作步骤") or ""),
            "expected_result": str(item.get("expected_result") or item.get("预期结果") or item.get("expected") or ""),
            "priority": priority if priority in {"high", "medium", "low"} else "medium",
            "test_type": test_type if test_type in {"functional", "api", "ui", "integration", "performance", "security"} else "functional",
            "raw": item,
        }

    @staticmethod
    def _deduplicate_cases(cases, seen_keys):
        result = []
        for case in cases:
            key = (case["title"].strip(), case["steps"].strip(), case["expected_result"].strip())
            if key in seen_keys:
                continue
            seen_keys.add(key)
            result.append(case)
        return result

    @staticmethod
    def _stringify_steps(steps):
        if isinstance(steps, list):
            return "\n".join(str(step) for step in steps)
        return str(steps)

    @classmethod
    def _parse_markdown_table(cls, content):
        rows = [line.strip() for line in content.splitlines() if line.strip().startswith("|")]
        data_rows = [row for row in rows if not re.match(r"^\|?\s*-+", row.replace("|", " |"))]
        if len(data_rows) < 2:
            return []
        headers = [part.strip() for part in data_rows[0].strip("|").split("|")]
        cases = []
        for index, row in enumerate(data_rows[1:], start=1):
            values = [part.strip() for part in row.strip("|").split("|")]
            item = dict(zip(headers, values))
            cases.append(cls._normalize_case(item, index))
        return cases
